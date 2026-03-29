"""
ptp_parser.py — Парсер описаний реальных пожаров из документов Word.
═══════════════════════════════════════════════════════════════════════════════
Извлекает структурированные данные из .docx-файлов с описаниями
реальных пожаров на резервуарах для загрузки в симулятор САУР-ПСП.

Отличие от планов тушения (ПТП):
  ПТП — расчётные параметры ДО пожара (нормативы, типовые схемы).
  Описание реального пожара — хронология ПОСЛЕ пожара: что произошло,
  какие решения принял РТП, какие были сбои, как развивалась обстановка.

Для обучения агентов ОП ценны именно описания реальных пожаров, потому что
они содержат ТРАЕКТОРИИ РЕШЕНИЙ — последовательности (состояние → действие),
которые являются обучающими данными для IRL, клонирования поведения,
обучения по предпочтениям.

Парсер извлекает:
  - Параметры объекта (тип РВС, объём, диаметр, топливо, кровля)
  - Ранг пожара
  - Хронологию с решениями РТП (Ч+N: кто, что сделал, результат)
  - Силы и средства (техника, стволы, ПНС, личный состав)
  - Пенные атаки (число, исход, причины неудач)
  - Инциденты (розливы, вторичные очаги, отказы техники)
  - Итог (локализован, ликвидирован, время, потери)
═══════════════════════════════════════════════════════════════════════════════
"""
from __future__ import annotations

import os
import re
import json
from dataclasses import dataclass, field, asdict
from typing import List, Dict, Optional, Tuple

try:
    from docx import Document
except ImportError:
    Document = None


# ═══════════════════════════════════════════════════════════════════════════
# СТРУКТУРА ДАННЫХ СЦЕНАРИЯ
# ═══════════════════════════════════════════════════════════════════════════
@dataclass
class TimelineEvent:
    """Одно событие хронологии ПТП."""
    t_min: int              # время от начала (мин)
    time_label: str         # «Ч+10», «14:35» и т.д.
    description: str        # описание события
    category: str = "info"  # "info", "warn", "danger", "success"

    def to_dict(self) -> dict:
        return asdict(self)


@dataclass
class ParsedScenario:
    """Результат парсинга одного ПТП."""
    # Источник
    source_file: str = ""
    parse_quality: float = 0.0  # 0..1 — насколько полно удалось извлечь

    # Параметры объекта
    rvs_type: str = "РВС"           # тип резервуара
    rvs_volume_m3: float = 0.0       # объём (м³)
    rvs_diameter_m: float = 0.0      # диаметр (м)
    fuel_type: str = "нефтепродукт"  # тип топлива
    roof_type: str = "неизвестно"    # «плавающая», «конусная», «понтонная»
    roof_obstruction: float = 0.0    # препятствие крыши [0; 1]

    # Параметры пожара
    fire_rank: int = 0               # ранг пожара (1–5)
    initial_fire_area_m2: float = 0.0
    total_duration_min: int = 0
    localized: bool = False
    extinguished: bool = False

    # Хронология
    timeline: List[TimelineEvent] = field(default_factory=list)

    # Силы и средства
    n_ac: int = 0          # автоцистерн
    n_apt: int = 0         # автопеноподъёмников
    n_pns: int = 0         # ПНС
    n_als: int = 0         # автолестниц / АКП
    n_total_personnel: int = 0  # общая численность ЛС

    # Пенные атаки
    foam_attacks_total: int = 0
    foam_attacks_successful: int = 0
    foam_attack_details: List[Dict] = field(default_factory=list)
    # [{"t_min": 340, "result": "неудача", "reason": "выход из строя ППП"}, ...]
    foam_type: str = ""    # тип пенообразователя

    # Инциденты (розливы, вторичные очаги, отказы)
    incidents: List[Dict] = field(default_factory=list)
    # [{"t_min": 557, "type": "розлив", "description": "300 м²"}, ...]

    # Решения РТП (для ОП: траектория действий эксперта)
    rtp_decisions: List[Dict] = field(default_factory=list)
    # [{"t_min": 12, "rtp": "РТП-1", "action": "создан штаб", "phase": "S2"}, ...]

    # Потери
    casualties_personnel: int = 0   # пострадавшие среди ЛС
    casualties_civilian: int = 0    # пострадавшие гражданские
    material_damage: str = ""       # описание материального ущерба

    # Нормативные данные (если указаны)
    foam_intensity_norm: float = 0.0    # л/(с·м²)
    cooling_flow_norm: float = 0.0      # л/с

    def to_dict(self) -> dict:
        d = asdict(self)
        d["timeline"] = [e.to_dict() for e in self.timeline]
        return d

    def to_simulator_format(self) -> dict:
        """Конвертировать в формат SCENARIOS из tank_fire_sim.py."""
        tl_lookup = {}
        timeline_list = []
        for ev in self.timeline:
            color_map = {"info": "#2980b9", "warn": "#e67e22",
                         "danger": "#c0392b", "success": "#27ae60"}
            color = color_map.get(ev.category, "#2980b9")
            entry = (ev.t_min, ev.time_label, ev.description, color)
            timeline_list.append(entry)
            tl_lookup.setdefault(ev.t_min, []).append(entry)

        return {
            "name": (f"{self.rvs_type} (V={self.rvs_volume_m3:.0f} м³, "
                     f"{self.fuel_type}, ранг №{self.fire_rank})"),
            "short": f"{self.rvs_type}-{self.rvs_volume_m3:.0f} (ранг {self.fire_rank})",
            "total_min": self.total_duration_min or 300,
            "initial_fire_area": self.initial_fire_area_m2 or 200.0,
            "fuel": self.fuel_type,
            "rvs_name": f"{self.rvs_type} (V={self.rvs_volume_m3:.0f} м³)",
            "rvs_diameter_m": self.rvs_diameter_m or 20.0,
            "fire_rank_default": self.fire_rank or 2,
            "roof_obstruction_init": self.roof_obstruction,
            "foam_intensity": self.foam_intensity_norm or 0.05,
            "tl_lookup": tl_lookup,
            "timeline": timeline_list,
            "scripted_effects": {},
            "actions_by_phase": None,
            "source_file": self.source_file,
        }


# ═══════════════════════════════════════════════════════════════════════════
# ПАТТЕРНЫ ПОИСКА
# ═══════════════════════════════════════════════════════════════════════════
_RE_VOLUME = re.compile(
    r'(?:V\s*=?\s*|объ[её]м[а-я]*\s*[=:—–-]?\s*)'
    r'(\d[\d\s]*\d|\d+)\s*м[³3]', re.IGNORECASE)

_RE_DIAMETER = re.compile(
    r'(?:D\s*=?\s*|диаметр[а-я]*\s*[=:—–-]?\s*)'
    r'(\d+[.,]?\d*)\s*м\b', re.IGNORECASE)

_RE_AREA = re.compile(
    r'(?:S\s*=?\s*|площад[ьи]\s*(?:горения|зеркала|пожара)?\s*[=:—–-]?\s*)'
    r'(\d[\d\s]*\d|\d+)\s*м[²2]', re.IGNORECASE)

_RE_RANK = re.compile(
    r'(?:ранг[а-я]*\s*(?:пожара)?\s*[=:—–-]?\s*[№#]?\s*)'
    r'(\d)', re.IGNORECASE)

_RE_TIME_HPLUS = re.compile(
    r'[ЧH]\s*\+\s*(\d+)', re.IGNORECASE)

_RE_TIME_HHMM = re.compile(
    r'(\d{1,2})\s*[:.ч]\s*(\d{2})\s*(?:мин)?')

_RE_FUEL = re.compile(
    r'(?:бензин|нефть|дизел|мазут|керосин|газоконденсат|нефтепродукт)',
    re.IGNORECASE)

_RE_ROOF = re.compile(
    r'(?:плавающ|понтон|конус|стацион|купол)',
    re.IGNORECASE)

_RE_FOAM_ATTACK = re.compile(
    r'(?:пенн(?:ая|ой)\s*атак|подач[аи]\s*пен)',
    re.IGNORECASE)

_RE_LOCALIZED = re.compile(
    r'(?:локализован|локализация)',
    re.IGNORECASE)

_RE_EXTINGUISHED = re.compile(
    r'(?:ликвидирован|потушен|горение\s*(?:прекращено|отсутствует))',
    re.IGNORECASE)

_RE_EQUIPMENT = {
    "АЦ": re.compile(r'(?:АЦ|автоцистерн)[\s-]*(\d+)?', re.IGNORECASE),
    "АПТ": re.compile(r'(?:АПТ|автопеноподъ[её]мник)[\s-]*(\d+)?', re.IGNORECASE),
    "ПНС": re.compile(r'(?:ПНС|насосн(?:ая|ой)\s*станц)[\s-]*(\d+)?', re.IGNORECASE),
    "АЛ": re.compile(r'(?:АЛ|АКП|автолестн|автоколенч)[\s-]*(\d+)?', re.IGNORECASE),
}

# Паттерны для описаний РЕАЛЬНЫХ пожаров (не ПТП)
_RE_RTP = re.compile(
    r'РТП[-\s]*(\d+)?|руководител[ья]\s*тушения', re.IGNORECASE)

_RE_FOAM_FAIL = re.compile(
    r'(?:прекращен|неудач|отказ|выход\s*из\s*строя|разрушен)',
    re.IGNORECASE)

_RE_FOAM_SUCCESS = re.compile(
    r'(?:горение\s*(?:отсутствует|прекратил)|ликвидирован[оа]?\s*горен|'
    r'успешн)',
    re.IGNORECASE)

_RE_INCIDENT = re.compile(
    r'(?:свищ|розлив|выброс|взрыв|обрушен|вторичн|загоран|воспламен|'
    r'отказ|авари)',
    re.IGNORECASE)

_RE_SHTAB = re.compile(
    r'(?:штаб|ОШ|оперативн\w*\s*штаб)', re.IGNORECASE)

_RE_BU = re.compile(
    r'(?:БУ|боев\w*\s*участ)', re.IGNORECASE)

_RE_STVOL = re.compile(
    r'(?:ствол|Антенор|Муссон|ГПС|ЛС-С|лафетн|РС-\d)', re.IGNORECASE)

_RE_DECISION_KEYWORDS = re.compile(
    r'(?:принял\s*руководств|назначен|создан|установлен|подан[ыо]?\s*ствол|'
    r'направлен|вызван|запрошен|объявлен|сформирован|развёрнут|'
    r'принял\s*решен|организован)',
    re.IGNORECASE)


# ═══════════════════════════════════════════════════════════════════════════
# ПАРСЕР
# ═══════════════════════════════════════════════════════════════════════════
class PTPParser:
    """Парсер плана тушения пожара из файла Word (.docx).

    Стратегия: извлекать всё, что удаётся найти по паттернам.
    Каждый найденный параметр повышает parse_quality.
    Если параметр не найден — используется значение по умолчанию.
    """

    def __init__(self):
        self.warnings: List[str] = []

    def parse(self, docx_path: str) -> ParsedScenario:
        """Распарсить один .docx файл ПТП.

        Возвращает ParsedScenario с заполненными полями.
        """
        if Document is None:
            raise ImportError("Установите python-docx: pip install python-docx")

        self.warnings = []
        result = ParsedScenario(source_file=os.path.basename(docx_path))

        doc = Document(docx_path)
        all_text = self._extract_all_text(doc)
        tables = self._extract_tables(doc)

        quality_points = 0
        max_points = 10

        # ── Объём РВС ─────────────────────────────────────────────────────
        m = _RE_VOLUME.search(all_text)
        if m:
            result.rvs_volume_m3 = float(m.group(1).replace(" ", ""))
            quality_points += 1
        else:
            self.warnings.append("Объём РВС не найден")

        # ── Диаметр ───────────────────────────────────────────────────────
        m = _RE_DIAMETER.search(all_text)
        if m:
            result.rvs_diameter_m = float(m.group(1).replace(",", "."))
            quality_points += 1
        elif result.rvs_volume_m3 > 0:
            # Оценка по объёму: D ≈ 2·√(V/(π·H)), H ≈ 12м для типового РВС
            import math
            result.rvs_diameter_m = round(
                2 * math.sqrt(result.rvs_volume_m3 / (math.pi * 12)), 1)
            quality_points += 0.5

        # ── Площадь пожара ────────────────────────────────────────────────
        m = _RE_AREA.search(all_text)
        if m:
            result.initial_fire_area_m2 = float(m.group(1).replace(" ", ""))
            quality_points += 1
        elif result.rvs_diameter_m > 0:
            import math
            result.initial_fire_area_m2 = round(
                math.pi * (result.rvs_diameter_m / 2) ** 2, 0)
            quality_points += 0.5

        # ── Ранг пожара ───────────────────────────────────────────────────
        m = _RE_RANK.search(all_text)
        if m:
            result.fire_rank = int(m.group(1))
            quality_points += 1
        else:
            # Оценка по объёму
            v = result.rvs_volume_m3
            if v >= 10000:
                result.fire_rank = 4
            elif v >= 3000:
                result.fire_rank = 3
            elif v >= 1000:
                result.fire_rank = 2
            else:
                result.fire_rank = 1
            quality_points += 0.3

        # ── Тип топлива ───────────────────────────────────────────────────
        m = _RE_FUEL.search(all_text)
        if m:
            fuel = m.group(0).lower()
            fuel_map = {"бензин": "бензин", "нефть": "нефть",
                        "дизел": "дизель", "мазут": "мазут",
                        "керосин": "керосин", "газоконденсат": "газоконденсат",
                        "нефтепродукт": "нефтепродукт"}
            for k, v in fuel_map.items():
                if k in fuel:
                    result.fuel_type = v
                    break
            quality_points += 1

        # ── Тип кровли ────────────────────────────────────────────────────
        m = _RE_ROOF.search(all_text)
        if m:
            roof = m.group(0).lower()
            if "плавающ" in roof or "понтон" in roof:
                result.roof_type = "плавающая"
                result.roof_obstruction = 0.70
            elif "конус" in roof or "купол" in roof:
                result.roof_type = "конусная"
                result.roof_obstruction = 0.0
            else:
                result.roof_type = "стационарная"
                result.roof_obstruction = 0.30
            quality_points += 1

        # ── Хронология ────────────────────────────────────────────────────
        timeline = self._parse_timeline(all_text, tables)
        result.timeline = timeline
        if len(timeline) >= 3:
            quality_points += 1
            result.total_duration_min = max(
                (ev.t_min for ev in timeline), default=300)

        # ── Пенные атаки с деталями ───────────────────────────────────────
        result.foam_attack_details = self._parse_foam_attacks(all_text, timeline)
        result.foam_attacks_total = len(result.foam_attack_details)
        result.foam_attacks_successful = sum(
            1 for fa in result.foam_attack_details if fa["result"] == "успех")
        if result.foam_attacks_total > 0:
            quality_points += 1

        # ── Инциденты (розливы, отказы, вторичные очаги) ───────────────────
        result.incidents = self._parse_incidents(timeline)
        if result.incidents:
            quality_points += 0.5

        # ── Решения РТП (траектория действий эксперта) ─────────────────────
        result.rtp_decisions = self._parse_decisions(timeline)
        if len(result.rtp_decisions) >= 3:
            quality_points += 1

        # ── Локализация и ликвидация ──────────────────────────────────────
        if _RE_LOCALIZED.search(all_text):
            result.localized = True
            quality_points += 0.5
        if _RE_EXTINGUISHED.search(all_text):
            result.extinguished = True
            quality_points += 0.5

        # ── Техника ───────────────────────────────────────────────────────
        for eq_name, pattern in _RE_EQUIPMENT.items():
            matches = pattern.findall(all_text)
            count = len(matches)
            if eq_name == "АЦ":
                result.n_ac = count
            elif eq_name == "АПТ":
                result.n_apt = count
            elif eq_name == "ПНС":
                result.n_pns = count
            elif eq_name == "АЛ":
                result.n_als = count

        # ── Нормативные данные ────────────────────────────────────────────
        foam_norms = {"бензин": 0.065, "нефть": 0.060,
                      "дизель": 0.048, "мазут": 0.060,
                      "керосин": 0.050, "нефтепродукт": 0.050}
        result.foam_intensity_norm = foam_norms.get(result.fuel_type, 0.050)

        result.parse_quality = min(1.0, quality_points / max_points)
        return result

    def _parse_foam_attacks(self, text: str,
                            timeline: List[TimelineEvent]) -> List[Dict]:
        """Извлечь детали пенных атак с результатами."""
        attacks = []
        for ev in timeline:
            if not _RE_FOAM_ATTACK.search(ev.description):
                continue
            result = "неизвестно"
            reason = ""
            # Искать исход в текущем и следующих событиях
            desc_lower = ev.description.lower()
            if _RE_FOAM_SUCCESS.search(desc_lower):
                result = "успех"
            elif _RE_FOAM_FAIL.search(desc_lower):
                result = "неудача"
                # Причина неудачи
                if "каркас" in desc_lower or "крыш" in desc_lower:
                    reason = "препятствие каркаса крыши"
                elif "выход из строя" in desc_lower or "отказ" in desc_lower:
                    reason = "отказ техники"
                elif "разрушен" in desc_lower:
                    reason = "разрушение пены"
                else:
                    reason = "не указана"
            attacks.append({
                "t_min": ev.t_min,
                "description": ev.description,
                "result": result,
                "reason": reason,
            })

        # Проверить следующие события на результат атаки
        for i, atk in enumerate(attacks):
            if atk["result"] == "неизвестно":
                # Искать в следующих событиях хронологии
                for ev in timeline:
                    if ev.t_min > atk["t_min"] and ev.t_min <= atk["t_min"] + 30:
                        if _RE_FOAM_FAIL.search(ev.description):
                            attacks[i]["result"] = "неудача"
                            break
                        elif _RE_FOAM_SUCCESS.search(ev.description):
                            attacks[i]["result"] = "успех"
                            break

        return attacks

    def _parse_incidents(self, timeline: List[TimelineEvent]) -> List[Dict]:
        """Извлечь инциденты (розливы, отказы, вторичные очаги)."""
        incidents = []
        for ev in timeline:
            if _RE_INCIDENT.search(ev.description):
                inc_type = "прочее"
                lower = ev.description.lower()
                if "розлив" in lower or "свищ" in lower:
                    inc_type = "розлив"
                elif "вторичн" in lower or "загоран" in lower or "воспламен" in lower:
                    inc_type = "вторичный очаг"
                elif "отказ" in lower or "выход из строя" in lower or "авари" in lower:
                    inc_type = "отказ техники"
                elif "выброс" in lower or "взрыв" in lower:
                    inc_type = "выброс/взрыв"
                elif "обрушен" in lower:
                    inc_type = "обрушение"
                incidents.append({
                    "t_min": ev.t_min,
                    "type": inc_type,
                    "description": ev.description,
                })
        return incidents

    def _parse_decisions(self, timeline: List[TimelineEvent]) -> List[Dict]:
        """Извлечь решения РТП из хронологии (траектория эксперта для ОП)."""
        decisions = []
        current_rtp = "РТП-1"

        for ev in timeline:
            # Смена РТП
            m_rtp = _RE_RTP.search(ev.description)
            if m_rtp and "принял" in ev.description.lower():
                num = m_rtp.group(1)
                if num:
                    current_rtp = f"РТП-{num}"

            # Проверить: содержит ли событие решение
            if not _RE_DECISION_KEYWORDS.search(ev.description):
                continue

            # Определить тип действия (маппинг на 15 действий симулятора)
            action_code = self._classify_action(ev.description)

            # Определить фазу по времени (эвристика)
            phase = self._estimate_phase(ev.t_min)

            decisions.append({
                "t_min": ev.t_min,
                "rtp": current_rtp,
                "description": ev.description,
                "action_code": action_code,
                "phase": phase,
            })

        return decisions

    @staticmethod
    def _classify_action(desc: str) -> str:
        """Классифицировать описание события → код действия (С1..О6)."""
        lower = desc.lower()
        if any(w in lower for w in ["спасен", "эвакуац"]):
            return "С1"
        if any(w in lower for w in ["защит", "охлажд"] + ["соседн"]):
            if "соседн" in lower:
                return "О2"
            return "С2"
        if "локализ" in lower:
            return "С3"
        if any(w in lower for w in ["пенн", "подач", "пен"]) and "атак" in lower:
            return "О3"
        if "вскипан" in lower or "выброс" in lower:
            return "С5"
        if any(w in lower for w in ["штаб", "ОШ", "боев"] + ["участ"]):
            if "участ" in lower or "БУ" in lower:
                return "Т1"
            return "Т1"
        if "перегруп" in lower or "переназнач" in lower:
            return "Т2"
        if any(w in lower for w in ["запрос", "вызов", "вызван", "направлен"]):
            return "Т3"
        if any(w in lower for w in ["ПНС", "насосн", "водоисточн", "установлен"]):
            return "Т4"
        if any(w in lower for w in ["ствол", "подан", "Антенор", "лафетн"]):
            return "О1"
        if any(w in lower for w in ["разведк", "уточн", "обследован"]):
            return "О4"
        if "розлив" in lower and ("ликвид" in lower or "тушен" in lower):
            return "О5"
        if any(w in lower for w in ["отход", "вывод", "отступ"]):
            return "О6"
        return "О4"  # по умолчанию — разведка

    @staticmethod
    def _estimate_phase(t_min: int) -> str:
        """Эвристическая оценка фазы пожара по времени."""
        if t_min <= 5:
            return "S1"
        elif t_min <= 30:
            return "S2"
        elif t_min <= 300:
            return "S3"
        elif t_min <= 600:
            return "S4"
        else:
            return "S5"

    def _extract_all_text(self, doc) -> str:
        """Собрать весь текст документа (абзацы + таблицы)."""
        parts = []
        for para in doc.paragraphs:
            parts.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    parts.append(cell.text)
        return "\n".join(parts)

    def _extract_tables(self, doc) -> List[List[List[str]]]:
        """Извлечь все таблицы как списки строк."""
        tables = []
        for table in doc.tables:
            rows = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows.append(cells)
            tables.append(rows)
        return tables

    def _parse_timeline(self, text: str,
                        tables: List) -> List[TimelineEvent]:
        """Извлечь хронологию событий."""
        events = []
        lines = text.split("\n")

        for line in lines:
            line = line.strip()
            if not line or len(line) < 10:
                continue

            t_min = None
            time_label = ""

            # Попытка: Ч+N
            m = _RE_TIME_HPLUS.search(line)
            if m:
                t_min = int(m.group(1))
                time_label = f"Ч+{t_min}"

            # Попытка: HH:MM (если есть базовое время)
            if t_min is None:
                m = _RE_TIME_HHMM.search(line)
                if m and any(kw in line.lower() for kw in
                             ["прибыт", "подан", "атак", "ликвидир",
                              "локализ", "обнаруж", "вызов", "создан",
                              "развёрт", "направл", "сообщен", "горен"]):
                    # Нет базового времени — пропускаем абсолютные часы
                    pass

            if t_min is None:
                continue

            # Определить категорию
            category = "info"
            lower = line.lower()
            if any(w in lower for w in ["опасн", "угроз", "крити",
                                         "свищ", "розлив", "взрыв",
                                         "прекращен", "неудач", "выход из строя"]):
                category = "danger"
            elif any(w in lower for w in ["атак", "готовн", "запрос"]):
                category = "warn"
            elif any(w in lower for w in ["ликвидир", "локализ",
                                           "потушен", "отсутств"]):
                category = "success"

            # Очистить описание от временно́й метки
            desc = line
            if time_label:
                desc = desc.replace(f"[{time_label}]", "").strip()
                desc = re.sub(r'^[ЧH]\s*\+\s*\d+\s*[:\-—–]?\s*', '', desc).strip()

            if desc:
                events.append(TimelineEvent(
                    t_min=t_min, time_label=time_label,
                    description=desc, category=category))

        # Сортировка по времени
        events.sort(key=lambda e: e.t_min)

        # Убрать дубликаты по времени
        seen = set()
        unique = []
        for ev in events:
            key = (ev.t_min, ev.description[:30])
            if key not in seen:
                seen.add(key)
                unique.append(ev)

        return unique


# ═══════════════════════════════════════════════════════════════════════════
# УТИЛИТЫ
# ═══════════════════════════════════════════════════════════════════════════
def parse_single(path: str) -> ParsedScenario:
    """Распарсить один файл ПТП."""
    parser = PTPParser()
    result = parser.parse(path)
    if parser.warnings:
        print(f"  Предупреждения ({os.path.basename(path)}):")
        for w in parser.warnings:
            print(f"    - {w}")
    return result


def save_scenario_json(scenario: ParsedScenario, output_dir: str) -> str:
    """Сохранить сценарий в JSON."""
    os.makedirs(output_dir, exist_ok=True)
    name = os.path.splitext(scenario.source_file)[0]
    name = re.sub(r'[^\w\-]', '_', name)
    path = os.path.join(output_dir, f"{name}.json")
    with open(path, "w", encoding="utf-8") as f:
        json.dump(scenario.to_dict(), f, ensure_ascii=False, indent=2,
                  default=str)
    return path
