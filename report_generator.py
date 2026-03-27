"""
report_generator.py
════════════════════════════════════════════════════════════════════════════════
Генератор отчётов по результатам моделирования тушения пожара РВС.

Форматы вывода:
  1. PDF-отчёт          — полный отчёт для оперативного штаба / руководства
  2. JSON-выгрузка      — структурированные данные для написания научной статьи
  3. DOCX-выгрузка      — текстовое резюме в формате Word (python-docx)

Зависимости: reportlab, python-docx (уже установлены)
════════════════════════════════════════════════════════════════════════════════
"""
from __future__ import annotations
import json
import math
import os
import io
import datetime
from typing import TYPE_CHECKING, List, Dict, Optional

# Импорт reportlab для генерации PDF
from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from reportlab.lib.units import mm, cm
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT, TA_JUSTIFY
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
    HRFlowable, PageBreak, Image as RLImage, KeepTogether,
)
from reportlab.platypus.flowables import Flowable
from reportlab.pdfgen import canvas as pdf_canvas
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

# Matplotlib для встраивания графиков в PDF
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.gridspec as gridspec
import numpy as np

if TYPE_CHECKING:
    from .tank_fire_sim import TankFireSim

# ══════════════════════════════════════════════════════════════════════════════
# ШРИФТЫ (поддержка кириллицы в reportlab)
# ══════════════════════════════════════════════════════════════════════════════

def _register_fonts():
    """Зарегистрировать шрифты с поддержкой кириллицы."""
    # Путь к шрифтам в Windows
    font_paths = [
        ("C:/Windows/Fonts/arial.ttf",     "Arial"),
        ("C:/Windows/Fonts/arialbd.ttf",   "Arial-Bold"),
        ("C:/Windows/Fonts/ariali.ttf",    "Arial-Italic"),
        ("C:/Windows/Fonts/cour.ttf",      "Courier-Cyr"),
        ("C:/Windows/Fonts/courbd.ttf",    "Courier-Cyr-Bold"),
    ]
    registered = {}
    for path, name in font_paths:
        if os.path.exists(path):
            try:
                pdfmetrics.registerFont(TTFont(name, path))
                registered[name] = True
            except Exception:
                pass
    return registered


_FONTS = _register_fonts()
_FONT     = "Arial"      if "Arial"      in _FONTS else "Helvetica"
_FONT_B   = "Arial-Bold" if "Arial-Bold" in _FONTS else "Helvetica-Bold"
_FONT_I   = "Arial-Italic"if "Arial-Italic" in _FONTS else "Helvetica-Oblique"
_FONT_MON = "Courier-Cyr" if "Courier-Cyr" in _FONTS else "Courier"


# ══════════════════════════════════════════════════════════════════════════════
# ЦВЕТА ОТЧЁТА
# ══════════════════════════════════════════════════════════════════════════════

C_FIRE    = colors.HexColor("#c0392b")   # красный — огонь / опасность
C_WATER   = colors.HexColor("#2471a3")   # синий — вода / охлаждение
C_SUCCESS = colors.HexColor("#27ae60")   # зелёный — успех / ликвидация
C_WARN    = colors.HexColor("#e67e22")   # оранжевый — предупреждение
C_ACCENT  = colors.HexColor("#1a5276")   # синий — фирменный акцент (шапка отчёта)
C_DARK    = colors.HexColor("#1a5276")   # alias → тот же акцентный синий
C_PANEL   = colors.HexColor("#eaf2fb")   # светло-синий фон панелей
C_HEADER  = colors.HexColor("#2c3e50")   # тёмно-серый — шапки таблиц
C_STRIPE  = colors.HexColor("#f4f6f7")   # полоска чередования строк
C_BORDER  = colors.HexColor("#aab7c4")   # рамки таблиц
C_CHART_BG= colors.HexColor("#f8f9fa")   # фон графиков (светлый)


# ══════════════════════════════════════════════════════════════════════════════
# ВСПОМОГАТЕЛЬНЫЕ ФУНКЦИИ
# ══════════════════════════════════════════════════════════════════════════════

def _para(text: str, style: ParagraphStyle) -> Paragraph:
    """Создать параграф с указанным стилем."""
    return Paragraph(text, style)


def _hr() -> HRFlowable:
    """Горизонтальная разделительная линия."""
    return HRFlowable(width="100%", thickness=0.5, color=C_BORDER, spaceAfter=4)


def _table_style(header_rows: int = 1) -> TableStyle:
    """Стандартный стиль для таблиц отчёта."""
    cmds = [
        # Шапка таблицы
        ("BACKGROUND",   (0, 0), (-1, header_rows - 1), C_HEADER),
        ("TEXTCOLOR",    (0, 0), (-1, header_rows - 1), colors.white),
        ("FONTNAME",     (0, 0), (-1, header_rows - 1), _FONT_B),
        ("FONTSIZE",     (0, 0), (-1, header_rows - 1), 9),
        ("ALIGN",        (0, 0), (-1, header_rows - 1), "CENTER"),
        # Данные
        ("FONTNAME",     (0, header_rows), (-1, -1), _FONT),
        ("FONTSIZE",     (0, header_rows), (-1, -1), 8),
        ("ALIGN",        (0, header_rows), (0, -1),  "LEFT"),
        ("ALIGN",        (1, header_rows), (-1, -1), "CENTER"),
        ("ROWBACKGROUNDS", (0, header_rows), (-1, -1), [colors.white, C_STRIPE]),
        # Рамки
        ("GRID",         (0, 0), (-1, -1), 0.4, C_BORDER),
        ("TOPPADDING",   (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING",(0, 0), (-1, -1), 4),
        ("LEFTPADDING",  (0, 0), (-1, -1), 6),
    ]
    return TableStyle(cmds)


def _fmt_time(t: int, base_h: int = 3, base_m: int = 20) -> str:
    """Перевести минуты симуляции в метку реального времени."""
    total_m = base_m + t
    h = (base_h + total_m // 60) % 24
    m = total_m % 60
    day = 14 + (base_h * 60 + base_m + t) // (24 * 60)
    return f"{day:02d}.03  {h:02d}:{m:02d}"


def _duration_str(minutes: int) -> str:
    """Форматировать продолжительность в ч мин."""
    h = minutes // 60
    m = minutes % 60
    if h > 0:
        return f"{h} ч {m} мин"
    return f"{m} мин"


# ══════════════════════════════════════════════════════════════════════════════
# ГЕНЕРАТОР ГРАФИКОВ ДЛЯ ОТЧЁТА
# ══════════════════════════════════════════════════════════════════════════════

def _build_metrics_figure(sim: "TankFireSim") -> io.BytesIO:
    """Построить фигуру метрик (4 субграфика) и вернуть PNG в памяти."""
    fig, axes = plt.subplots(2, 2, figsize=(14, 8),
                             facecolor="#ffffff",
                             gridspec_kw=dict(hspace=0.52, wspace=0.38))
    fig.patch.set_facecolor("#ffffff")

    ts_f  = [x[0] for x in sim.h_fire]
    val_f = [x[1] for x in sim.h_fire]
    ts_w  = [x[0] for x in sim.h_water]
    val_w = [x[1] for x in sim.h_water]
    ts_r  = [x[0] for x in sim.h_risk]
    val_r = [x[1] for x in sim.h_risk]
    ts_t  = [x[0] for x in sim.h_trunks]
    val_t = [x[1] for x in sim.h_trunks]

    panel_bg = "#f8f9fa"
    text_c   = "#2c3e50"
    grid_c   = "#dee2e6"
    tick_c   = "#5d6d7e"

    def _style_ax(ax, title, ylabel=""):
        ax.set_facecolor(panel_bg)
        ax.set_title(title, color=text_c, fontsize=10, pad=5, fontweight="bold")
        ax.tick_params(colors=tick_c, labelsize=8)
        ax.set_xlabel("Время, мин", color=tick_c, fontsize=8)
        if ylabel:
            ax.set_ylabel(ylabel, color=tick_c, fontsize=8)
        for sp in ax.spines.values():
            sp.set_color(grid_c)
            sp.set_linewidth(0.8)
        ax.grid(True, color=grid_c, linewidth=0.6, alpha=0.9, linestyle="--")

    # Площадь пожара
    ax = axes[0, 0]
    _style_ax(ax, "Рис. 1а. Площадь пожара S(t)", "м²")
    if ts_f:
        ax.plot(ts_f, val_f, color="#c0392b", linewidth=2.0, label="S(t)", zorder=3)
        ax.fill_between(ts_f, val_f, alpha=0.18, color="#c0392b")
        ax.set_xlim(0, max(ts_f[-1], 100))
    ax.legend(fontsize=8, facecolor="white", edgecolor=grid_c, labelcolor=text_c,
              framealpha=0.9)

    # Расход ОВ
    ax = axes[0, 1]
    _style_ax(ax, "Рис. 1б. Расход ОВ Q(t)", "л/с")
    if ts_w:
        ax.plot(ts_w, val_w, color="#1a5276", linewidth=2.0, label="Q(t)", zorder=3)
        ax.fill_between(ts_w, val_w, alpha=0.15, color="#1a5276")
        ax.set_xlim(0, max(ts_w[-1], 100))
    ax.legend(fontsize=8, facecolor="white", edgecolor=grid_c, labelcolor=text_c,
              framealpha=0.9)

    # Число стволов
    ax = axes[1, 0]
    _style_ax(ax, "Рис. 1в. Число стволов N(t)", "ед.")
    if ts_t:
        ax.step(ts_t, val_t, color="#2471a3", linewidth=2.0, where="post",
                label="N стволов", zorder=3)
        ax.fill_between(ts_t, val_t, alpha=0.12, color="#2471a3", step="post")
        ax.axhline(7, color="#27ae60", linewidth=1.2, linestyle="--", alpha=0.8,
                   label="цель: 7 стволов")
        ax.set_xlim(0, max(ts_t[-1], 100))
    ax.legend(fontsize=8, facecolor="white", edgecolor=grid_c, labelcolor=text_c,
              framealpha=0.9)

    # Индекс риска
    ax = axes[1, 1]
    _style_ax(ax, "Рис. 1г. Индекс риска R(t)", "безразм.")
    if ts_r:
        ax.plot(ts_r, val_r, color="#c0392b", linewidth=2.0, label="R(t)", zorder=3)
        ax.fill_between(ts_r, val_r, alpha=0.12, color="#c0392b")
        ax.axhline(0.75, color="#c0392b", linewidth=1.2, linestyle="--",
                   alpha=0.8, label="критич. порог (0.75)")
        ax.axhline(0.50, color="#e67e22", linewidth=1.0, linestyle=":",
                   alpha=0.7, label="высок. порог (0.50)")
        ax.set_ylim(0, 1.05)
        ax.set_xlim(0, max(ts_r[-1], 100))
    ax.legend(fontsize=7, facecolor="white", edgecolor=grid_c, labelcolor=text_c,
              framealpha=0.9)

    fig.suptitle("Динамика ключевых показателей тушения пожара РВС",
                 color=text_c, fontsize=11, fontweight="bold", y=1.01)
    plt.tight_layout()
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=150, bbox_inches="tight",
                facecolor="#ffffff")
    buf.seek(0)
    plt.close(fig)
    return buf


def _build_rl_figure(sim: "TankFireSim") -> io.BytesIO:
    """Построить фигуру RL-агента (Q-значения + частота выбора + кривая наград)."""
    try:
        from .tank_fire_sim import ACTIONS, LEVEL_C, N_ACT, P
    except ImportError:
        from tank_fire_sim import ACTIONS, LEVEL_C, N_ACT, P

    fig, axes = plt.subplots(1, 3, figsize=(14, 4.5),
                             facecolor="#ffffff",
                             gridspec_kw=dict(wspace=0.38))
    fig.patch.set_facecolor("#ffffff")
    panel_bg = "#f8f9fa"
    text_c   = "#2c3e50"
    grid_c   = "#dee2e6"
    tick_c   = "#5d6d7e"

    def _style_ax(ax, title):
        ax.set_facecolor(panel_bg)
        ax.set_title(title, color=text_c, fontsize=9, pad=5, fontweight="bold")
        ax.tick_params(colors=tick_c, labelsize=7)
        for sp in ax.spines.values():
            sp.set_color(grid_c)
            sp.set_linewidth(0.8)
        ax.grid(True, color=grid_c, linewidth=0.5, alpha=0.9, linestyle="--")

    codes = [a[0] for a in ACTIONS]
    cols  = [LEVEL_C[a[1]] for a in ACTIONS]

    # Q-значения
    ax = axes[0]
    _style_ax(ax, "Рис. 2а. Q-значения действий")
    qv = sim.agent.q_values(sim._state())
    bars = ax.bar(range(N_ACT), qv, color=cols, alpha=0.85, width=0.7)
    if 0 <= sim.last_action < N_ACT:
        bars[sim.last_action].set_edgecolor("#c0392b")
        bars[sim.last_action].set_linewidth(2.0)
    ax.set_xticks(range(N_ACT))
    ax.set_xticklabels(codes, rotation=45, ha="right", fontsize=7, color=tick_c)
    ax.set_ylabel("Q-ценность", color=tick_c, fontsize=8)

    # Частота выбора
    ax = axes[1]
    _style_ax(ax, "Рис. 2б. Частота выбора действий")
    cnt = sim.agent.action_counts
    if cnt.sum() > 0:
        ax.bar(range(N_ACT), cnt / max(cnt.sum(), 1), color=cols, alpha=0.8, width=0.7)
    ax.set_xticks(range(N_ACT))
    ax.set_xticklabels(codes, rotation=45, ha="right", fontsize=7, color=tick_c)
    ax.set_ylabel("Доля", color=tick_c, fontsize=8)

    # Кривая наград
    ax = axes[2]
    _style_ax(ax, "Рис. 2в. Накопленная награда")
    if sim.h_reward:
        rw = sim.h_reward
        cumrew = np.cumsum(rw)
        ax.plot(cumrew, color="#27ae60", linewidth=1.8, label="Sigma reward", zorder=3)
        if len(rw) > 20:
            ma = np.convolve(rw, np.ones(20) / 20, mode="valid")
            ax.plot(range(19, len(rw)), ma.cumsum() +
                    (cumrew[18] if len(cumrew) > 18 else 0),
                    color="#e67e22", linewidth=1.2, alpha=0.85, label="MA-20")
    ax.legend(fontsize=7, facecolor="white", edgecolor=grid_c, labelcolor=text_c,
              framealpha=0.9)
    ax.set_ylabel("Sigma r", color=tick_c, fontsize=8)

    fig.suptitle("Анализ RL-агента: Q-значения, частота действий, динамика наград",
                 color=text_c, fontsize=10, fontweight="bold")
    plt.tight_layout()
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=150, bbox_inches="tight",
                facecolor="#ffffff")
    buf.seek(0)
    plt.close(fig)
    return buf


# ══════════════════════════════════════════════════════════════════════════════
# СБОР СТАТИСТИКИ ПО ИТОГАМ СИМУЛЯЦИИ
# ══════════════════════════════════════════════════════════════════════════════

def collect_stats(sim: "TankFireSim") -> dict:
    """Собрать итоговую статистику из объекта TankFireSim.

    Returns:
        dict с полями: scenario, duration_min, final_phase, fire_area_final,
                        max_fire_area, min_fire_area, total_water_ls_avg,
                        foam_attacks, extinguished, localized,
                        max_risk, mean_risk, n_trunks_max,
                        rl_epsilon, rl_total_reward, action_distribution,
                        events_count, timeline_events
    """
    h_fire  = sim.h_fire
    h_water = sim.h_water
    h_risk  = sim.h_risk
    h_trunks= sim.h_trunks

    fire_vals  = [v for _, v in h_fire]
    water_vals = [v for _, v in h_water]
    risk_vals  = [v for _, v in h_risk]
    trunk_vals = [v for _, v in h_trunks]

    try:
        from .tank_fire_sim import ACTIONS, SCENARIOS
    except ImportError:
        from tank_fire_sim import ACTIONS, SCENARIOS
    # Для пользовательских сценариев cfg может быть в sim._cfg
    cfg = SCENARIOS.get(sim.scenario) or getattr(sim, "_cfg", {})
    # Гарантировать наличие обязательных полей
    cfg.setdefault("name", sim.scenario)
    cfg.setdefault("short", cfg["name"][:20])
    cfg.setdefault("rvs_name", "РВС")
    cfg.setdefault("fuel", "бензин")
    cfg.setdefault("fire_rank_default", 2)
    cfg.setdefault("total_min", 300)
    cfg.setdefault("initial_fire_area", 168.0)

    # Распределение действий
    ac = sim.agent.action_counts
    total_ac = max(ac.sum(), 1)
    action_dist = {
        ACTIONS[i][0]: {
            "count": int(ac[i]),
            "fraction": float(ac[i] / total_ac),
            "description": ACTIONS[i][2],
        }
        for i in range(len(ACTIONS))
        if ac[i] > 0
    }
    action_dist_sorted = dict(
        sorted(action_dist.items(), key=lambda x: x[1]["count"], reverse=True)
    )

    # События из журнала
    events_log = [
        {"t_min": t, "color": c, "text": txt}
        for t, c, txt in sim.events[:200]   # ограничить 200 записями
    ]

    return {
        # Сценарий
        "scenario_key":    sim.scenario,
        "scenario_name":   cfg["name"],
        "rvs_name":        cfg["rvs_name"],
        "fuel":            cfg["fuel"],
        "fire_rank":       cfg["fire_rank_default"],
        "total_sim_min":   cfg["total_min"],
        # Время
        "sim_duration_min": sim.t,
        "sim_duration_str": _duration_str(sim.t),
        "generated_at":    datetime.datetime.now().isoformat(timespec="seconds"),
        # Фаза и исход
        "final_phase":     sim.phase,
        "extinguished":    sim.extinguished,
        "localized":       sim.localized,
        "outcome":         "ликвидирован" if sim.extinguished else
                           ("локализован" if sim.localized else "активный"),
        # Площадь пожара
        "fire_area_initial_m2": cfg["initial_fire_area"],
        "fire_area_final_m2":   sim.fire_area,
        "fire_area_max_m2":     max(fire_vals) if fire_vals else 0,
        "fire_area_min_m2":     min(fire_vals) if fire_vals else 0,
        # Водоснабжение
        "water_flow_max_ls":    max(water_vals) if water_vals else 0,
        "water_flow_mean_ls":   sum(water_vals) / max(len(water_vals), 1),
        # Пенные атаки
        "foam_attacks_total":   sim.foam_attacks,
        "foam_flow_last_ls":    sim.foam_flow_ls,
        "roof_obstruction":     sim.roof_obstruction,
        "akp50_used":           sim.akp50_available,
        # Силы
        "trunks_burn_final":    sim.n_trunks_burn,
        "trunks_burn_max":      max(trunk_vals) if trunk_vals else 0,
        "trunks_nbr_final":     sim.n_trunks_nbr,
        "n_pns":                sim.n_pns,
        "n_bu":                 sim.n_bu,
        "has_shtab":            sim.has_shtab,
        "secondary_fire":       sim.secondary_fire,
        "spill":                sim.spill,
        # Риск
        "risk_max":             max(risk_vals) if risk_vals else 0,
        "risk_mean":            sum(risk_vals) / max(len(risk_vals), 1),
        "risk_final":           sim._risk(),
        # RL-агент
        "rl_epsilon":           sim.agent.epsilon,
        "rl_total_reward":      float(sum(sim.h_reward)) if sim.h_reward else 0.0,
        "rl_steps":             len(sim.h_reward),
        "rl_episodes":          len(sim.agent.episode_rewards),
        "action_distribution":  action_dist_sorted,
        # Хронология
        "events_count":         len(sim.events),
        "events_log":           events_log,
        # Нормативы ГОСТ
        "foam_intensity_norm":  cfg["foam_intensity"],
        "q_foam_required_ls":   cfg["foam_intensity"] * cfg["initial_fire_area"],
        "q_cooling_required_ls": 0.8 * math.pi * cfg["rvs_diameter_m"],
    }


# ══════════════════════════════════════════════════════════════════════════════
# PDF-ОТЧЁТ
# ══════════════════════════════════════════════════════════════════════════════

class _PageTemplate:
    """Колонтитулы страниц PDF-отчёта."""
    def __init__(self, title: str, stats: dict):
        self.title = title
        self.stats = stats

    def __call__(self, canv: pdf_canvas.Canvas, doc):
        canv.saveState()
        w, h = A4
        # Верхний колонтитул — синяя полоса
        canv.setFillColor(colors.HexColor("#1a5276"))
        canv.rect(0, h - 16*mm, w, 16*mm, fill=1, stroke=0)
        canv.setFont(_FONT_B, 9)
        canv.setFillColor(colors.white)
        canv.drawString(15*mm, h - 10*mm, "САУР ПСП — Отчёт о моделировании тушения пожара РВС")
        canv.setFont(_FONT, 8)
        canv.drawRightString(w - 15*mm, h - 10*mm,
                             f"Создан: {self.stats['generated_at']}")
        # Тонкая линия под колонтитулом
        canv.setStrokeColor(colors.HexColor("#aab7c4"))
        canv.setLineWidth(0.5)
        canv.line(15*mm, h - 16.5*mm, w - 15*mm, h - 16.5*mm)
        # Нижний колонтитул — светло-серая полоса
        canv.setFillColor(colors.HexColor("#eaf2fb"))
        canv.rect(0, 0, w, 11*mm, fill=1, stroke=0)
        canv.setStrokeColor(colors.HexColor("#aab7c4"))
        canv.line(15*mm, 11*mm, w - 15*mm, 11*mm)
        canv.setFont(_FONT, 7)
        canv.setFillColor(colors.HexColor("#2c3e50"))
        canv.drawString(15*mm, 3.5*mm,
                        "САУР ПСП v1.0 — Система адаптивного управления реагированием пожарно-спасательного подразделения")
        canv.drawRightString(w - 15*mm, 3.5*mm, f"Страница {doc.page}")
        canv.restoreState()


def generate_pdf_report(sim: "TankFireSim", output_path: str = "") -> str:
    """Сгенерировать PDF-отчёт по результатам моделирования.

    Args:
        sim:         экземпляр TankFireSim после завершения симуляции
        output_path: путь к выходному файлу (если пусто — авто-имя в текущей папке)

    Returns:
        Путь к созданному PDF-файлу
    """
    stats = collect_stats(sim)

    # Авто-имя файла
    if not output_path:
        ts = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        output_path = os.path.join(
            os.path.dirname(__file__),
            f"report_{stats['scenario_key']}_{ts}.pdf"
        )

    # Стили текста
    base_styles = getSampleStyleSheet()
    sty = {
        "title": ParagraphStyle("title", fontName=_FONT_B, fontSize=20,
                                textColor=colors.white, alignment=TA_CENTER,
                                spaceAfter=6),
        "subtitle": ParagraphStyle("subtitle", fontName=_FONT, fontSize=11,
                                   textColor=C_PANEL, alignment=TA_CENTER,
                                   spaceAfter=4),
        "h1": ParagraphStyle("h1", fontName=_FONT_B, fontSize=13,
                             textColor=C_FIRE, spaceBefore=12, spaceAfter=4),
        "h2": ParagraphStyle("h2", fontName=_FONT_B, fontSize=10,
                             textColor=C_HEADER, spaceBefore=8, spaceAfter=3),
        "body": ParagraphStyle("body", fontName=_FONT, fontSize=9,
                               textColor=colors.black, leading=14,
                               spaceAfter=4, alignment=TA_JUSTIFY),
        "mono": ParagraphStyle("mono", fontName=_FONT_MON, fontSize=8,
                               textColor=C_HEADER, spaceAfter=2),
        "note": ParagraphStyle("note", fontName=_FONT_I, fontSize=8,
                               textColor=colors.gray, spaceAfter=3),
        "ok": ParagraphStyle("ok", fontName=_FONT_B, fontSize=9,
                             textColor=C_SUCCESS, spaceAfter=3),
        "warn": ParagraphStyle("warn", fontName=_FONT_B, fontSize=9,
                               textColor=C_WARN, spaceAfter=3),
    }

    # Инициализация документа
    doc = SimpleDocTemplate(
        output_path,
        pagesize=A4,
        leftMargin=20*mm, rightMargin=20*mm,
        topMargin=25*mm, bottomMargin=20*mm,
        title=f"Отчёт о моделировании — {stats['scenario_name']}",
        author="САУР ПСП v1.0",
        subject="Тушение пожара РВС",
    )

    page_cb = _PageTemplate(stats["scenario_name"], stats)
    story: list = []

    # ═══════════════════════════════════════════════════════════
    # ТИТУЛЬНАЯ СТРАНИЦА
    # ═══════════════════════════════════════════════════════════
    story.append(Spacer(1, 35*mm))

    # Акцентная линия + синий блок-заголовок
    story.append(HRFlowable(width="100%", thickness=4,
                             color=colors.HexColor("#c0392b"), spaceAfter=4))
    cover_data = [[
        Paragraph("ОТЧЁТ О МОДЕЛИРОВАНИИ<br/>УПРАВЛЕНИЯ ТУШЕНИЕМ ПОЖАРА РВС", sty["title"])
    ]]
    cover_tbl = Table(cover_data, colWidths=[170*mm])
    cover_tbl.setStyle(TableStyle([
        ("BACKGROUND",    (0, 0), (-1, -1), colors.HexColor("#1a5276")),
        ("TOPPADDING",    (0, 0), (-1, -1), 16),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 16),
        ("LEFTPADDING",   (0, 0), (-1, -1), 14),
        ("RIGHTPADDING",  (0, 0), (-1, -1), 14),
    ]))
    story.append(cover_tbl)
    story.append(HRFlowable(width="100%", thickness=3,
                             color=colors.HexColor("#e67e22"), spaceAfter=6))
    story.append(Spacer(1, 5*mm))

    sub_sty = ParagraphStyle("pdf_sub", fontName=_FONT_B, fontSize=11,
                              textColor=colors.HexColor("#1a5276"),
                              alignment=TA_CENTER, spaceAfter=4)
    story.append(_para(stats["scenario_name"], sub_sty))
    story.append(Spacer(1, 4*mm))

    # Краткая сводка на обложке
    outcome_icon = "✅" if stats["extinguished"] else ("🔒" if stats["localized"] else "🔥")
    outcome_text = outcome_icon + " " + stats["outcome"].upper()
    outcome_style = sty["ok"] if stats["extinguished"] else (
        sty["warn"] if stats["localized"] else sty["warn"])
    story.append(Paragraph(f"<b>Исход:</b> {outcome_text}", outcome_style))
    story.append(Spacer(1, 6*mm))

    cover_meta = [
        ["Параметр", "Значение"],
        ["Сценарий",           stats["scenario_name"][:60]],
        ["Объект",             stats["rvs_name"]],
        ["Горючее",            stats["fuel"]],
        ["Ранг пожара",        f"№{stats['fire_rank']}"],
        ["Продолжительность",  stats["sim_duration_str"]],
        ["Шагов симуляции",    str(stats["rl_steps"])],
        ["Дата создания",      stats["generated_at"].replace("T", "  ")],
    ]
    ct = Table(cover_meta, colWidths=[70*mm, 100*mm])
    ct.setStyle(_table_style(1))
    story.append(ct)
    story.append(PageBreak())

    # ═══════════════════════════════════════════════════════════
    # 1. ИСХОДНЫЕ ДАННЫЕ И НОРМАТИВНЫЕ ТРЕБОВАНИЯ
    # ═══════════════════════════════════════════════════════════
    story.append(_para("1. Исходные данные и нормативные требования", sty["h1"]))
    story.append(_hr())

    story.append(_para(
        f"Моделирование выполнено на основе сценария «{stats['scenario_name']}» "
        f"с использованием физической модели пенного тушения по "
        f"ГОСТ Р 51043-2002 и СП 155.13130.2014.",
        sty["body"]
    ))
    story.append(Spacer(1, 3*mm))

    norms_data = [
        ["Параметр", "Значение", "Норматив"],
        ["Площадь зеркала горения",
         f"{stats['fire_area_initial_m2']:.0f} м²",
         "S = π·D²/4"],
        ["Норм. интенсивность подачи пены",
         f"{stats['foam_intensity_norm']:.3f} л/(м²·с)",
         "ГОСТ Р 51043 / СП 155"],
        ["Требуемый расход пенного раствора",
         f"{stats['q_foam_required_ls']:.1f} л/с",
         "Q_пена = I_норм × S"],
        ["Требуемый расход охлаждения",
         f"{stats['q_cooling_required_ls']:.1f} л/с",
         "q = 0.8 л/(с·м) × π·D"],
        ["Препятствие каркаса крыши",
         f"{stats['roof_obstruction']*100:.0f}%",
         "Плавающая крыша / конус"],
    ]
    nt = Table(norms_data, colWidths=[80*mm, 50*mm, 40*mm])
    nt.setStyle(_table_style(1))
    story.append(nt)
    story.append(Spacer(1, 5*mm))

    # ═══════════════════════════════════════════════════════════
    # 2. ИТОГИ МОДЕЛИРОВАНИЯ
    # ═══════════════════════════════════════════════════════════
    story.append(_para("2. Итоги моделирования", sty["h1"]))
    story.append(_hr())

    results_data = [
        ["Показатель", "Значение"],
        ["Исход пожара",          stats["outcome"].title()],
        ["Продолжительность",     stats["sim_duration_str"]],
        ["Финальная фаза",        stats["final_phase"]],
        ["Площадь пожара (нач.)", f"{stats['fire_area_initial_m2']:.0f} м²"],
        ["Площадь пожара (макс.)",f"{stats['fire_area_max_m2']:.0f} м²"],
        ["Площадь пожара (фин.)", f"{stats['fire_area_final_m2']:.0f} м²"],
        ["Пенных атак проведено", str(stats["foam_attacks_total"])],
        ["АКП-50 задействован",   "Да" if stats["akp50_used"] else "Нет"],
        ["Стволов охлаждения (макс.)", str(stats["trunks_burn_max"])],
        ["ПНС/ПАНРК на воде",    str(stats["n_pns"])],
        ["Боевых участков",      str(stats["n_bu"])],
        ["Вторичный пожар",      "Был" if stats["secondary_fire"] else "Нет"],
        ["Розлив горящего топлива","Был" if stats["spill"] else "Нет"],
        ["Расход ОВ (макс.)",    f"{stats['water_flow_max_ls']:.0f} л/с"],
        ["Индекс риска (макс.)", f"{stats['risk_max']:.2f}"],
        ["Индекс риска (сред.)", f"{stats['risk_mean']:.2f}"],
    ]
    rt = Table(results_data, colWidths=[95*mm, 75*mm])
    rt.setStyle(_table_style(1))
    story.append(rt)
    story.append(Spacer(1, 5*mm))

    # ═══════════════════════════════════════════════════════════
    # 3. ГРАФИКИ ДИНАМИКИ ПОКАЗАТЕЛЕЙ
    # ═══════════════════════════════════════════════════════════
    story.append(_para("3. Графики динамики показателей", sty["h1"]))
    story.append(_hr())
    story.append(_para(
        "На рисунке 1 представлены временны́е ряды ключевых показателей симуляции: "
        "площадь пожара, расход огнетушащих веществ, число стволов охлаждения "
        "и интегральный индекс риска.",
        sty["body"]
    ))
    story.append(Spacer(1, 3*mm))

    # Встроить граф метрик
    metrics_buf = _build_metrics_figure(sim)
    img_metrics = RLImage(metrics_buf, width=170*mm, height=97*mm)
    story.append(img_metrics)
    story.append(_para(
        "Рис. 1. Динамика ключевых показателей тушения пожара РВС (результаты моделирования).",
        sty["note"]
    ))
    story.append(Spacer(1, 5*mm))

    # RL-графики
    story.append(_para("4. Результаты RL-агента", sty["h1"]))
    story.append(_hr())
    story.append(_para(
        f"Управление тушением пожара осуществлялось Q-learning агентом "
        f"(ε-greedy, ε={stats['rl_epsilon']:.2f}). "
        f"За {stats['rl_steps']} шагов симуляции агент выполнил "
        f"{sum(v['count'] for v in stats['action_distribution'].values())} действий, "
        f"суммарная накопленная награда составила {stats['rl_total_reward']:.1f}.",
        sty["body"]
    ))
    story.append(Spacer(1, 3*mm))

    rl_buf = _build_rl_figure(sim)
    img_rl = RLImage(rl_buf, width=170*mm, height=49*mm)
    story.append(img_rl)
    story.append(_para(
        "Рис. 2. Q-значения действий, частота выбора и кривая накопленной награды RL-агента.",
        sty["note"]
    ))
    story.append(Spacer(1, 5*mm))

    # Таблица распределения действий
    story.append(_para("Распределение действий РТП (RL-агент)", sty["h2"]))
    act_hdr = [["Код", "Описание", "Кол-во", "Доля, %"]]
    act_rows = [
        [code,
         info["description"][:55],
         str(info["count"]),
         f"{info['fraction']*100:.1f}"]
        for code, info in list(stats["action_distribution"].items())[:15]
    ]
    act_data = act_hdr + act_rows
    at = Table(act_data, colWidths=[18*mm, 105*mm, 22*mm, 25*mm])
    at.setStyle(_table_style(1))
    story.append(at)
    story.append(PageBreak())

    # ═══════════════════════════════════════════════════════════
    # 5. ХРОНОЛОГИЯ СОБЫТИЙ
    # ═══════════════════════════════════════════════════════════
    story.append(_para("5. Журнал событий симуляции", sty["h1"]))
    story.append(_hr())
    story.append(_para(
        f"Ниже приведены первые 60 событий из журнала симуляции "
        f"(всего зафиксировано {stats['events_count']} событий).",
        sty["note"]
    ))
    story.append(Spacer(1, 3*mm))

    ev_hdr  = [["Время, мин", "Событие"]]
    ev_rows = []
    try:
        from .tank_fire_sim import P as _P
    except ImportError:
        from tank_fire_sim import P as _P
    _COLOR_TAG = {
        _P["danger"]:  "❗ ", _P["warn"]: "⚠ ",
        _P["success"]: "✅ ", _P["info"]: "ℹ ",
    }
    for ev in stats["events_log"][:60]:
        prefix = _COLOR_TAG.get(ev["color"], "   ")
        ev_rows.append([
            str(ev["t_min"]),
            (prefix + ev["text"])[:100],
        ])
    ev_data = ev_hdr + ev_rows
    ev_tbl = Table(ev_data, colWidths=[25*mm, 145*mm])
    ev_tbl.setStyle(_table_style(1))
    story.append(ev_tbl)
    story.append(Spacer(1, 5*mm))

    # ═══════════════════════════════════════════════════════════
    # 6. ВЫВОДЫ И РЕКОМЕНДАЦИИ
    # ═══════════════════════════════════════════════════════════
    story.append(_para("6. Выводы и рекомендации", sty["h1"]))
    story.append(_hr())

    concl_sty = ParagraphStyle("pdf_concl", fontName=_FONT, fontSize=9,
                               textColor=colors.HexColor("#2c3e50"), leading=14,
                               leftIndent=10, spaceAfter=5)
    conclusions = []
    if stats["extinguished"]:
        conclusions.append(
            "Пожар ликвидирован в ходе симуляции. RL-агент освоил стратегию управления "
            "ресурсами, обеспечив выполнение нормативных условий для успешной пенной атаки."
        )
    elif stats["localized"]:
        conclusions.append(
            "Пожар локализован, но не ликвидирован. Рекомендуется увеличить число эпизодов "
            "обучения RL-агента и расход пенообразователя."
        )
    else:
        conclusions.append(
            "Пожар не локализован. Требуется перепараметризация сценария: увеличить запас "
            "пенообразователя и задействовать АКП-50 на более раннем этапе."
        )

    conclusions.append(
        f"Проведено {stats['foam_attacks_total']} пенных атак. Многократные атаки "
        f"обусловлены препятствием каркаса крыши ({stats.get('roof_obstruction',0)*100:.0f}%), "
        f"характерным для РВС с плавающей крышей."
    )
    if stats["akp50_used"]:
        conclusions.append(
            "АКП-50 задействован — снижение препятствия каркаса до 20% является ключевым "
            "тактическим решением для обеспечения нормативного расхода пены."
        )
    if stats["risk_max"] > 0.75:
        conclusions.append(
            f"Максимальный индекс риска {stats['risk_max']:.2f} превысил критический порог. "
            f"Рекомендуется ранняя установка ПНС и наращивание водяного охлаждения."
        )
    conclusions.append(
        f"RL-агент (ε = {stats['rl_epsilon']:.2f}): для эксплуатационного режима "
        f"снизить ε до 0.05–0.10. Обучение на смешанном наборе сценариев "
        f"повышает устойчивость политики управления."
    )

    for i, txt in enumerate(conclusions, 1):
        numbered = f"<b><font color='#1a5276'>{i}.</font></b>  {txt}"
        story.append(_para(numbered, concl_sty))
        story.append(Spacer(1, 2*mm))

    story.append(Spacer(1, 5*mm))
    story.append(_para(
        "Отчёт сформирован автоматически системой САУР ПСП v1.0. "
        "Все расчёты выполнены в соответствии с ГОСТ Р 51043-2002 и СП 155.13130.2014.",
        sty["note"]
    ))

    # Сборка PDF
    doc.build(story, onFirstPage=page_cb, onLaterPages=page_cb)
    return output_path


# ══════════════════════════════════════════════════════════════════════════════
# РАСШИРЕННЫЙ PDF-ОТЧЁТ С FLAT vs HRL СРАВНЕНИЕМ
# ══════════════════════════════════════════════════════════════════════════════

def generate_comprehensive_report(sim: "TankFireSim",
                                   flat_sim=None,
                                   hier_sim=None,
                                   batch_data=None,
                                   output_path: str = "") -> str:
    """Сгенерировать расширенный PDF-отчёт.

    Args:
        sim:         основной экземпляр TankFireSim
        flat_sim:    обученный плоский агент (TankFireSim) для сравнения, или None
        hier_sim:    обученный иерархический агент (HierarchicalTankFireSim), или None
        batch_data:  результаты массового эксперимента (dict), или None
        output_path: путь к выходному файлу (если пусто — авто-имя)

    Returns:
        Путь к созданному PDF-файлу
    """
    stats = collect_stats(sim)

    if not output_path:
        ts = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        output_path = os.path.join(
            os.path.dirname(__file__),
            f"report_{stats['scenario_key']}_{ts}.pdf"
        )

    try:
        from .tank_fire_sim import ACTIONS, SCENARIOS
    except ImportError:
        from tank_fire_sim import ACTIONS, SCENARIOS

    # ── Стили (светлая тема) ─────────────────────────────────────────────────
    sty = {
        "title": ParagraphStyle("cr_title", fontName=_FONT_B, fontSize=18,
                                textColor=colors.white, alignment=TA_CENTER,
                                spaceAfter=6, leading=24),
        "subtitle": ParagraphStyle("cr_subtitle", fontName=_FONT_B, fontSize=11,
                                   textColor=colors.HexColor("#1a5276"),
                                   alignment=TA_CENTER, spaceAfter=4),
        "h1": ParagraphStyle("cr_h1", fontName=_FONT_B, fontSize=13,
                              textColor=colors.HexColor("#1a5276"),
                              spaceBefore=14, spaceAfter=5,
                              borderPad=3, borderColor=colors.HexColor("#1a5276"),
                              borderWidth=0),
        "h2": ParagraphStyle("cr_h2", fontName=_FONT_B, fontSize=10,
                              textColor=colors.HexColor("#2c3e50"),
                              spaceBefore=8, spaceAfter=3),
        "body": ParagraphStyle("cr_body", fontName=_FONT, fontSize=9,
                               textColor=colors.HexColor("#2c3e50"), leading=14,
                               spaceAfter=5, alignment=TA_JUSTIFY),
        "mono": ParagraphStyle("cr_mono", fontName=_FONT_MON, fontSize=8,
                               textColor=colors.HexColor("#1a5276"),
                               backColor=colors.HexColor("#eaf2fb"),
                               borderPad=4, spaceAfter=4, spaceBefore=2),
        "note": ParagraphStyle("cr_note", fontName=_FONT_I, fontSize=8,
                               textColor=colors.HexColor("#5d6d7e"),
                               spaceAfter=3),
        "ok":   ParagraphStyle("cr_ok",   fontName=_FONT_B, fontSize=9,
                               textColor=C_SUCCESS, spaceAfter=3),
        "warn": ParagraphStyle("cr_warn", fontName=_FONT_B, fontSize=9,
                               textColor=C_WARN, spaceAfter=3),
        "eq":   ParagraphStyle("cr_eq",   fontName=_FONT_MON, fontSize=9,
                               textColor=colors.HexColor("#1a5276"),
                               backColor=colors.HexColor("#eaf2fb"),
                               alignment=TA_CENTER, borderPad=6,
                               spaceAfter=5, spaceBefore=5),
        "concl": ParagraphStyle("cr_concl", fontName=_FONT, fontSize=9,
                                textColor=colors.HexColor("#2c3e50"), leading=14,
                                leftIndent=10, spaceAfter=5,
                                firstLineIndent=0),
    }

    doc = SimpleDocTemplate(
        output_path, pagesize=A4,
        leftMargin=20*mm, rightMargin=20*mm,
        topMargin=25*mm, bottomMargin=20*mm,
        title=f"Научный отчёт — {stats['scenario_name']}",
        author="САУР ПСП v1.0",
        subject="Сравнительное исследование RL-агентов управления тушением пожара РВС",
    )

    page_cb = _PageTemplate(stats["scenario_name"], stats)
    story: list = []

    # ═══════════════════════════════════════════════════════════
    # ТИТУЛЬНАЯ СТРАНИЦА
    # ═══════════════════════════════════════════════════════════
    story.append(Spacer(1, 20*mm))

    # Тонкая акцентная линия над заголовком
    story.append(HRFlowable(width="100%", thickness=4,
                             color=colors.HexColor("#c0392b"), spaceAfter=4))

    cover_data = [[Paragraph(
        "НАУЧНЫЙ ОТЧЁТ<br/>"
        "<font size=14>Сравнительное исследование алгоритмов управления<br/>"
        "тушением пожара резервуарного парка<br/>"
        "на основе методов обучения с подкреплением</font>",
        sty["title"]
    )]]
    cover_tbl = Table(cover_data, colWidths=[170*mm])
    cover_tbl.setStyle(TableStyle([
        ("BACKGROUND",    (0, 0), (-1, -1), colors.HexColor("#1a5276")),
        ("TOPPADDING",    (0, 0), (-1, -1), 18),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 18),
        ("LEFTPADDING",   (0, 0), (-1, -1), 14),
        ("RIGHTPADDING",  (0, 0), (-1, -1), 14),
    ]))
    story.append(cover_tbl)
    # Оранжевая полоса под заголовком
    story.append(HRFlowable(width="100%", thickness=3,
                             color=colors.HexColor("#e67e22"), spaceAfter=6))
    story.append(Spacer(1, 5*mm))

    # Подзаголовок сценария
    sub_style = ParagraphStyle("cover_sub", fontName=_FONT_B, fontSize=11,
                                textColor=colors.HexColor("#1a5276"),
                                alignment=TA_CENTER, spaceAfter=4)
    story.append(_para(stats["scenario_name"], sub_style))
    story.append(Spacer(1, 5*mm))

    outcome_icon = "✅" if stats["extinguished"] else ("🔒" if stats["localized"] else "🔥")
    outcome_sty = sty["ok"] if stats["extinguished"] else sty["warn"]
    story.append(_para(f"<b>Исход симуляции:</b> {outcome_icon} {stats['outcome'].upper()}",
                       outcome_sty))
    story.append(Spacer(1, 6*mm))

    cover_meta = [
        ["Параметр", "Значение"],
        ["Сценарий",            stats["scenario_name"][:60]],
        ["Объект защиты",       stats["rvs_name"]],
        ["Горючее",             stats["fuel"]],
        ["Ранг пожара",         f"№{stats['fire_rank']}"],
        ["Продолжительность",   stats["sim_duration_str"]],
        ["Шагов симуляции",     str(stats["rl_steps"])],
        ["Flat Q-learning",     "Есть" if flat_sim is not None else "Нет"],
        ["Иерархический RL",    "Есть" if hier_sim is not None else "Нет"],
        ["Дата создания",       stats["generated_at"].replace("T", "  ")],
    ]
    ct = Table(cover_meta, colWidths=[70*mm, 100*mm])
    ct.setStyle(_table_style(1))
    story.append(ct)
    story.append(PageBreak())

    # ═══════════════════════════════════════════════════════════
    # РАЗДЕЛ 1 — ОБЪЕКТ И СЦЕНАРИЙ
    # ═══════════════════════════════════════════════════════════
    story.append(_para("1. Объект исследования и сценарий моделирования", sty["h1"]))
    story.append(_hr())

    cfg = SCENARIOS.get(sim.scenario) or getattr(sim, "_cfg", {})
    story.append(_para(
        f"Объектом моделирования является резервуарный парк нефтепродуктов: "
        f"{stats['rvs_name']}, горючее — {stats['fuel']}, ранг пожара №{stats['fire_rank']}. "
        f"Физическая модель распространения пожара основана на законе Стефана–Больцмана "
        f"и нормативах СП 155.13130.2014 / ГОСТ Р 51043-2002.",
        sty["body"]
    ))
    story.append(Spacer(1, 3*mm))

    story.append(_para("Модель распространения пожара:", sty["h2"]))
    story.append(_para("dS/dt = v_spread × P(t)", sty["eq"]))
    story.append(_para(
        "где S(t) — площадь зеркала горения (м²), v_spread — скорость распространения (м²/мин), "
        "P(t) — периметр горения (м). Скорость v_spread зависит от вида горючего, "
        "условий ветра и задержки подачи пенного раствора.",
        sty["body"]
    ))
    story.append(Spacer(1, 2*mm))

    story.append(_para("Нормативные требования (ГОСТ Р 51043-2002 / СП 155.13130.2014):", sty["h2"]))
    norms_data = [
        ["Параметр", "Значение", "Норматив"],
        ["Площадь зеркала горения",      f"{stats['fire_area_initial_m2']:.0f} м²",   "S = π·D²/4"],
        ["Нормативная интенсивность пены",f"{stats['foam_intensity_norm']:.3f} л/(м²·с)", "ГОСТ Р 51043"],
        ["Требуемый расход пенного р-ра", f"{stats['q_foam_required_ls']:.1f} л/с",    "Q = I_норм × S"],
        ["Требуемый расход охлаждения",  f"{stats['q_cooling_required_ls']:.1f} л/с", "q = 0.8·π·D"],
        ["Препятствие каркаса крыши",    f"{stats['roof_obstruction']*100:.0f}%",      "Плав. крыша / конус"],
    ]
    nt = Table(norms_data, colWidths=[80*mm, 50*mm, 40*mm])
    nt.setStyle(_table_style(1))
    story.append(nt)
    story.append(Spacer(1, 5*mm))

    # ═══════════════════════════════════════════════════════════
    # РАЗДЕЛ 2 — МЕТОД УПРАВЛЕНИЯ
    # ═══════════════════════════════════════════════════════════
    story.append(_para("2. Метод управления: Q-learning и иерархический RL", sty["h1"]))
    story.append(_hr())

    story.append(_para("2.1 Плоский Q-learning агент", sty["h2"]))
    story.append(_para(
        f"Пространство состояний: s ∈ {{0..255}} (8-битная дискретизация физических переменных). "
        f"Пространство действий: |A| = {len(ACTIONS)} действий (S1–S5 стратегические, "
        f"T1–T4 тактические, O1–O6 оперативные).",
        sty["body"]
    ))
    story.append(_para("Уравнение Беллмана (Q-update):", sty["h2"]))
    story.append(_para(
        "Q(s,a) ← Q(s,a) + α[r + γ·max_a' Q(s',a') − Q(s,a)]",
        sty["eq"]
    ))
    story.append(_para(
        "где α — скорость обучения, γ — коэффициент дисконтирования, "
        "r — мгновенная награда, s' — следующее состояние.",
        sty["body"]
    ))
    story.append(Spacer(1, 3*mm))

    story.append(_para("2.2 Иерархический RL (3-уровневая архитектура)", sty["h2"]))
    story.append(_para(
        "Иерархический агент реализует вертикаль управления ICS (Incident Command System) "
        "согласно ГОСТ Р 22.7.01-2021:\n"
        "• L3 (НГ/ГУ МЧС): стратегический уровень — Q-таблица размером 32×3;\n"
        "• L2 (РТП/НШ): тактический уровень — Q-таблица размером 64×5;\n"
        "• L1 (НБТП): оперативный уровень — Q-таблица размером 256×15.",
        sty["body"]
    ))
    story.append(_para("Составная функция награды L1:", sty["h2"]))
    story.append(_para(
        "r_total = r_env + λ·r_intrinsic",
        sty["eq"]
    ))
    story.append(_para(
        "где r_intrinsic поощряет L1 за выбор действий, согласованных с макро-целью L2; "
        "λ — коэффициент баланса (настраивается через GUI).",
        sty["body"]
    ))
    story.append(Spacer(1, 5*mm))

    # ═══════════════════════════════════════════════════════════
    # РАЗДЕЛ 3 — ИТОГИ МОДЕЛИРОВАНИЯ
    # ═══════════════════════════════════════════════════════════
    story.append(_para("3. Итоги моделирования", sty["h1"]))
    story.append(_hr())

    results_data = [
        ["Показатель", "Значение"],
        ["Исход пожара",             stats["outcome"].title()],
        ["Продолжительность",        stats["sim_duration_str"]],
        ["Финальная фаза",           stats["final_phase"]],
        ["Площадь пожара — нач.",    f"{stats['fire_area_initial_m2']:.0f} м²"],
        ["Площадь пожара — макс.",   f"{stats['fire_area_max_m2']:.0f} м²"],
        ["Площадь пожара — фин.",    f"{stats['fire_area_final_m2']:.0f} м²"],
        ["Пенных атак",              str(stats["foam_attacks_total"])],
        ["АКП-50",                   "Задействован" if stats["akp50_used"] else "Нет"],
        ["Стволов охлаждения — макс.",str(stats["trunks_burn_max"])],
        ["ПНС/ПАНРК на воде",       str(stats["n_pns"])],
        ["Боевых участков",         str(stats["n_bu"])],
        ["Вторичный пожар",         "Был" if stats["secondary_fire"] else "Нет"],
        ["Розлив топлива",          "Был" if stats["spill"] else "Нет"],
        ["Расход ОВ — макс.",       f"{stats['water_flow_max_ls']:.0f} л/с"],
        ["Индекс риска — макс.",    f"{stats['risk_max']:.2f}"],
        ["Индекс риска — средн.",   f"{stats['risk_mean']:.2f}"],
    ]
    rt = Table(results_data, colWidths=[95*mm, 75*mm])
    rt.setStyle(_table_style(1))
    story.append(rt)
    story.append(Spacer(1, 5*mm))

    # ═══════════════════════════════════════════════════════════
    # РАЗДЕЛ 4 — ДИНАМИКА ПОКАЗАТЕЛЕЙ (ГРАФИКИ)
    # ═══════════════════════════════════════════════════════════
    story.append(_para("4. Динамика показателей моделирования", sty["h1"]))
    story.append(_hr())
    story.append(_para(
        "На рисунке 1 представлены временны́е ряды ключевых показателей симуляции: "
        "площадь пожара S(t), суммарный расход ОВ Q(t), число стволов охлаждения N(t) "
        "и интегральный индекс риска R(t).",
        sty["body"]
    ))
    story.append(Spacer(1, 3*mm))

    metrics_buf = _build_metrics_figure(sim)
    story.append(RLImage(metrics_buf, width=170*mm, height=97*mm))
    story.append(_para(
        "Рис. 1. Динамика ключевых показателей тушения пожара РВС.",
        sty["note"]
    ))
    story.append(Spacer(1, 5*mm))

    # ═══════════════════════════════════════════════════════════
    # РАЗДЕЛ 5 — АНАЛИЗ RL-АГЕНТА
    # ═══════════════════════════════════════════════════════════
    story.append(_para("5. Анализ RL-агента", sty["h1"]))
    story.append(_hr())
    story.append(_para(
        f"Управление осуществлялось Q-learning агентом (ε={stats['rl_epsilon']:.3f}). "
        f"За {stats['rl_steps']} шагов агент накопил суммарную награду "
        f"{stats['rl_total_reward']:.1f}.",
        sty["body"]
    ))
    story.append(Spacer(1, 3*mm))

    rl_buf = _build_rl_figure(sim)
    story.append(RLImage(rl_buf, width=170*mm, height=49*mm))
    story.append(_para(
        "Рис. 2. Q-значения действий, частота выбора и кривая накопленной награды RL-агента.",
        sty["note"]
    ))
    story.append(Spacer(1, 4*mm))

    story.append(_para("Топ-10 действий по частоте выбора:", sty["h2"]))
    act_hdr = [["Код", "Описание", "Кол-во", "Доля, %"]]
    act_rows = [
        [code,
         info["description"][:55],
         str(info["count"]),
         f"{info['fraction']*100:.1f}"]
        for code, info in list(stats["action_distribution"].items())[:10]
    ]
    at = Table(act_hdr + act_rows, colWidths=[18*mm, 105*mm, 22*mm, 25*mm])
    at.setStyle(_table_style(1))
    story.append(at)
    story.append(PageBreak())

    # ═══════════════════════════════════════════════════════════
    # РАЗДЕЛ 6 — СРАВНЕНИЕ FLAT vs ИЕРАРХИЧЕСКИЙ RL
    # ═══════════════════════════════════════════════════════════
    story.append(_para("6. Сравнение: Flat Q-learning vs Иерархический RL", sty["h1"]))
    story.append(_hr())

    if flat_sim is not None or hier_sim is not None:
        def _safe_stat(s, attr, default=None):
            try:
                return getattr(s, attr, default)
            except Exception:
                return default

        cmp_rows = [["Метрика", "Flat Q-learning", "Иерархический RL", "Δ"]]

        # Collect stats from flat_sim
        flat_stats = collect_stats(flat_sim) if flat_sim is not None else None
        hier_stats = None
        if hier_sim is not None:
            try:
                env = (getattr(hier_sim, "env", None)
                       or getattr(hier_sim, "sim", None)
                       or getattr(hier_sim, "_env", None))
                if env is not None:
                    hier_stats = collect_stats(env)
                    # Добавить специфические метрики иерархического агента
                    try:
                        hier_stats["hrl_l1_coverage"] = f"{hier_sim.l1.coverage():.0%}"
                        hier_stats["hrl_l2_coverage"] = f"{hier_sim.l2.coverage():.0%}"
                        hier_stats["hrl_l3_coverage"] = f"{hier_sim.l3.coverage():.0%}"
                        rw_l1 = hier_sim.l1.episode_rewards
                        if rw_l1:
                            hier_stats["rl_total_reward"] = sum(rw_l1[-50:]) / max(len(rw_l1[-50:]),1)
                    except Exception:
                        pass
            except Exception:
                pass

        def _fmt_val(s, key, fallback="—"):
            if s is None:
                return fallback
            v = s.get(key)
            if v is None:
                return fallback
            if isinstance(v, float):
                return f"{v:.2f}"
            return str(v)

        def _delta(s1, s2, key, higher_better=True):
            if s1 is None or s2 is None:
                return "—"
            v1 = s1.get(key)
            v2 = s2.get(key)
            if v1 is None or v2 is None:
                return "—"
            try:
                d = float(v2) - float(v1)
                pct = d / max(abs(float(v1)), 1e-9) * 100
                sign = "+" if d >= 0 else ""
                return f"{sign}{pct:.1f}%"
            except Exception:
                return "—"

        metrics_cmp = [
            ("Исход",             "outcome",          False),
            ("Длительность",      "sim_duration_str", False),
            ("Площадь — финал. м²","fire_area_final_m2", False),
            ("Пенных атак",       "foam_attacks_total", False),
            ("Σ Награда (агент)", "rl_total_reward",  True),
            ("Риск — макс.",      "risk_max",         False),
            ("Риск — средн.",     "risk_mean",        False),
            ("Расход ОВ, л/с",   "water_flow_max_ls", False),
            ("АКП-50 применён",  "akp50_used",        True),
            ("Покрытие L1 (HRL)","hrl_l1_coverage",  True),
            ("Покрытие L2 (HRL)","hrl_l2_coverage",  True),
            ("Покрытие L3 (HRL)","hrl_l3_coverage",  True),
        ]
        for label, key, hb in metrics_cmp:
            cmp_rows.append([
                label,
                _fmt_val(flat_stats, key),
                _fmt_val(hier_stats, key),
                _delta(flat_stats, hier_stats, key, hb),
            ])

        cmp_tbl = Table(cmp_rows, colWidths=[55*mm, 38*mm, 47*mm, 30*mm])
        cmp_tbl.setStyle(_table_style(1))
        story.append(cmp_tbl)
        story.append(Spacer(1, 4*mm))

        story.append(_para(
            "Интерпретация: положительное Δ в колонке 'Δ' означает улучшение "
            "иерархического агента относительно плоского по данной метрике. "
            "Для метрик риска и длительности меньшее значение предпочтительнее.",
            sty["note"]
        ))
    else:
        story.append(_para(
            "Данные для сравнения недоступны. Обучите оба агента (Flat Q-learning и "
            "Иерархический RL) на вкладке «Иерархический RL» и повторно сформируйте отчёт.",
            sty["warn"]
        ))
    story.append(Spacer(1, 5*mm))

    # ═══════════════════════════════════════════════════════════
    # РАЗДЕЛ 7 — НОРМАТИВНЫЙ АНАЛИЗ
    # ═══════════════════════════════════════════════════════════
    story.append(_para("7. Нормативный анализ (ГОСТ Р 51043-2002 / СП 155.13130.2014)", sty["h1"]))
    story.append(_hr())

    q_foam_req  = stats["q_foam_required_ls"]
    q_foam_act  = stats.get("foam_flow_last_ls", 0)
    q_cool_req  = stats["q_cooling_required_ls"]
    q_cool_act  = stats.get("water_flow_mean_ls", 0)
    roof_obs    = stats.get("roof_obstruction", 0)
    foam_eff    = q_foam_act * (1 - roof_obs)

    def _norm_status(actual, required, desc=""):
        if required <= 0:
            return "—"
        ok = actual >= required * 0.95
        return "✅ Выполнено" if ok else "❌ Не выполнено"

    norm_data = [
        ["Требование", "Требуемое", "Фактическое", "Выполнение"],
        ["Расход пенного р-ра",
         f"{q_foam_req:.1f} л/с",
         f"{q_foam_act:.1f} л/с",
         _norm_status(q_foam_act, q_foam_req)],
        ["Расход охлаждения горящего РВС",
         f"{q_cool_req:.1f} л/с",
         f"{q_cool_act:.1f} л/с",
         _norm_status(q_cool_act, q_cool_req)],
        ["Эффективный расход пены (с учётом крыши)",
         f"{q_foam_req:.1f} л/с",
         f"{foam_eff:.1f} л/с",
         _norm_status(foam_eff, q_foam_req)],
        ["Препятствие каркаса крыши",
         "≤ 30%",
         f"{roof_obs*100:.0f}%",
         "✅ Норма" if roof_obs <= 0.30 else "⚠ Превышение"],
        ["Число пенных атак",
         "≥ 1",
         str(stats["foam_attacks_total"]),
         "✅ Выполнено" if stats["foam_attacks_total"] >= 1 else "❌ Нет"],
    ]
    nrm_tbl = Table(norm_data, colWidths=[70*mm, 35*mm, 35*mm, 30*mm])
    nrm_tbl.setStyle(_table_style(1))
    story.append(nrm_tbl)
    story.append(Spacer(1, 5*mm))

    # ═══════════════════════════════════════════════════════════
    # РАЗДЕЛ 8 — ХРОНОЛОГИЯ СОБЫТИЙ
    # ═══════════════════════════════════════════════════════════
    story.append(_para("8. Хронология событий симуляции", sty["h1"]))
    story.append(_hr())
    story.append(_para(
        f"Приведены первые 50 событий из журнала симуляции "
        f"(всего зафиксировано {stats['events_count']} событий).",
        sty["note"]
    ))
    story.append(Spacer(1, 3*mm))

    try:
        from .tank_fire_sim import P as _P
    except ImportError:
        from tank_fire_sim import P as _P
    _COLOR_TAG = {
        _P["danger"]:  "❗ ", _P["warn"]: "⚠ ",
        _P["success"]: "✅ ", _P["info"]: "ℹ ",
    }

    ev_hdr  = [["Время, мин", "Событие"]]
    ev_rows = []
    for ev in stats["events_log"][:50]:
        prefix = _COLOR_TAG.get(ev["color"], "   ")
        ev_rows.append([
            str(ev["t_min"]),
            (prefix + ev["text"])[:100],
        ])
    ev_tbl = Table(ev_hdr + ev_rows, colWidths=[25*mm, 145*mm])
    ev_tbl.setStyle(_table_style(1))
    story.append(ev_tbl)
    story.append(Spacer(1, 5*mm))

    # ═══════════════════════════════════════════════════════════
    # РАЗДЕЛ 9 — ВЫВОДЫ И РЕКОМЕНДАЦИИ
    # ═══════════════════════════════════════════════════════════
    story.append(_para("9. Выводы и рекомендации", sty["h1"]))
    story.append(_hr())
    story.append(_para(
        "На основании результатов моделирования сформулированы следующие выводы:",
        sty["body"]
    ))
    story.append(Spacer(1, 3*mm))

    conclusions = []

    # ── Вспомогательные значения ──────────────────────────────────────────
    dur_str   = stats.get("sim_duration_str", "—")
    dur_min   = stats.get("sim_duration_min", 0)
    fire_init = stats.get("fire_area_initial_m2", 0)
    fire_fin  = stats.get("fire_area_final_m2", 0)
    fire_max  = stats.get("fire_area_max_m2", 0)
    risk_max  = stats.get("risk_max", 0)
    risk_mean = stats.get("risk_mean", 0)
    foam_n    = stats.get("foam_attacks_total", 0)
    roof_obs  = stats.get("roof_obstruction", 0)
    eps_fin   = stats.get("rl_epsilon", 0)
    steps_n   = stats.get("rl_steps", 0)
    rew_tot   = stats.get("rl_total_reward", 0)
    ep_rews   = stats.get("rl_episode_rewards", [])
    n_ep      = len(ep_rews) if ep_rews else 0
    akp50     = stats.get("akp50_used", False)
    water_max = stats.get("water_flow_max_ls", 0)
    water_avg = stats.get("water_flow_mean_ls", 0)
    scenario  = stats.get("scenario_name", "—")
    rank      = stats.get("fire_rank", "—")

    # ── Вывод 1 — Хронология и исход пожара ──────────────────────────────
    if stats["extinguished"]:
        outcome_txt = (
            f"Пожар <b>успешно ликвидирован</b> в сценарии «{scenario}» (ранг {rank}). "
            f"Общее время симуляции составило <b>{dur_str}</b> ({dur_min} мин). "
            f"Площадь горения сократилась с {fire_init:.0f} м² до 0 м² "
            f"(максимум в ходе пожара: {fire_max:.0f} м²). "
            f"Результат подтверждает способность RL-агента обеспечить ликвидацию "
            f"в рамках нормативного времени реагирования МЧС."
        )
    elif stats["localized"]:
        outcome_txt = (
            f"Пожар <b>локализован</b>, но полная ликвидация не достигнута "
            f"(сценарий «{scenario}», ранг {rank}, время: {dur_str}). "
            f"Конечная площадь горения — {fire_fin:.0f} м² (макс. {fire_max:.0f} м²). "
            f"Рекомендуется увеличить число эпизодов обучения, повысить запас "
            f"пенообразователя и снизить ε до 0.05 для эксплуатационного режима."
        )
    else:
        outcome_txt = (
            f"Пожар <b>не был локализован</b> за {dur_str} симуляции "
            f"(сценарий «{scenario}», ранг {rank}). "
            f"Площадь горения достигла максимума {fire_max:.0f} м². "
            f"Требуется перепараметризация: увеличить запас пенообразователя, "
            f"задействовать АКП-50 раньше, увеличить число ПНС."
        )
    conclusions.append(outcome_txt)

    # ── Вывод 2 — Водоснабжение и ресурсы ────────────────────────────────
    conclusions.append(
        f"Максимальный расход огнетушащего вещества составил <b>{water_max:.0f} л/с</b> "
        f"(средний: {water_avg:.0f} л/с). "
        f"Всего проведено <b>{foam_n} пенных атак</b>. "
        f"Препятствие каркаса плавающей крыши: {roof_obs*100:.0f}% — "
        + (
            "АКП-50 <b>применён</b>: препятствие снижено до 20%, что обеспечило "
            "нормативный расход пены по ГОСТ Р 51043-2002."
            if akp50 else
            "АКП-50 <b>не применялся</b>: при препятствии крыши ≥70% рекомендуется "
            "раннее задействование коленчатого подъёмника для достижения нормативного "
            "расхода пены."
        )
    )

    # ── Вывод 3 — Уровень риска ───────────────────────────────────────────
    risk_level = "критический" if risk_max > 0.75 else ("повышенный" if risk_max > 0.5 else "умеренный")
    risk_advice = (
        "Для снижения риска рекомендуется: ранняя установка ПНС (Ч+5 мин), "
        "немедленное охлаждение соседнего РВС, эвакуация персонала из зоны 50 м."
        if risk_max > 0.75 else
        "Своевременная подача стволов и установка ПНС обеспечили контролируемый "
        "уровень угрозы на протяжении большей части симуляции."
    )
    conclusions.append(
        f"Максимальный интегральный индекс риска: <b>{risk_max:.2f}</b> "
        f"(уровень — <b>{risk_level}</b>; критический порог: 0.75). "
        f"Средний индекс риска по симуляции: {risk_mean:.2f}. "
        f"{risk_advice}"
    )

    # ── Вывод 4 — Flat Q-learning агент: итерации и эффективность ────────
    train_note = (
        f"Агент обучался в течение <b>{n_ep} эпизодов</b> ({steps_n} шагов)."
        if n_ep > 0 else
        "Обучение агента не проводилось (режим inference или ручное управление)."
    )
    conv_note = (
        f"Финальный ε = {eps_fin:.3f} свидетельствует о "
        + ("высокой степени эксплуатации — агент преимущественно выбирает лучшие действия."
           if eps_fin < 0.15 else
           "продолжающемся исследовании — рекомендуется дообучение для снижения ε до 0.05.")
    )
    conclusions.append(
        f"<b>Flat Q-learning агент.</b> {train_note} "
        f"Суммарная накопленная награда: <b>{rew_tot:.1f}</b>. {conv_note} "
        f"<u>Преимущества Flat RL:</u> простота реализации, низкая вычислительная стоимость, "
        f"прозрачность Q-таблицы. "
        f"<u>Недостатки:</u> единственный уровень абстракции не отражает вертикаль команды МЧС, "
        f"сложность масштабирования при расширении пространства состояний, "
        f"медленная сходимость при большом числе действий."
    )

    # ── Вывод 5 — Иерархический RL: сравнение с Flat ──────────────────────
    if hier_sim is not None and flat_sim is not None:
        try:
            flat_env = flat_sim
            hier_env = (getattr(hier_sim, "env", None)
                        or getattr(hier_sim, "sim", None)
                        or getattr(hier_sim, "_env", None))
            _f_out = "ликвидирован" if flat_env.extinguished else ("локализован" if flat_env.localized else "активный")
            _h_out = "ликвидирован" if hier_env.extinguished else ("локализован" if hier_env.localized else "активный")
            _f_t   = flat_env.t
            _h_t   = getattr(hier_env, "t", "—")
            _f_r   = sum(flat_env.agent.episode_rewards[-20:]) / max(len(flat_env.agent.episode_rewards[-20:]), 1) if flat_env.agent.episode_rewards else 0
            cov_l1 = f"{hier_sim.l1.coverage():.0%}" if hasattr(hier_sim, "l1") else "—"
            cov_l2 = f"{hier_sim.l2.coverage():.0%}" if hasattr(hier_sim, "l2") else "—"
            cov_l3 = f"{hier_sim.l3.coverage():.0%}" if hasattr(hier_sim, "l3") else "—"
            # Determine scenario labels for each agent
            cur_scen   = getattr(sim,      "scenario", "текущий")
            flat_scen  = getattr(flat_env, "scenario", None)
            hier_scen  = getattr(hier_env, "scenario", None) or getattr(hier_sim, "scenario", None)
            flat_label = f" [сценарий: {flat_scen}]" if flat_scen and flat_scen != cur_scen else ""
            hier_label = f" [сценарий: {hier_scen}]" if hier_scen and hier_scen != cur_scen else ""
            cross_note = ""
            if (flat_scen and flat_scen != cur_scen) or (hier_scen and hier_scen != cur_scen):
                cross_note = (
                    f"<i>Примечание: данные обучения агентов получены на сценарии обучения "
                    f"({flat_scen or hier_scen}), тогда как текущая симуляция выполнена на "
                    f"сценарии «{cur_scen}». Время обучения (Flat: {_f_t} мин, "
                    f"HRL: {_h_t} мин) не совпадает со временем текущей симуляции — "
                    f"это нормально и отражает сложность обучающего сценария.</i> "
                )
            hrl_detail = (
                f"{cross_note}"
                f"Flat агент{flat_label}: исход = {_f_out}, время обучения = {_f_t} мин, "
                f"средняя награда (посл. 20 эп.) = {_f_r:.1f}. "
                f"HRL агент (L3/L2/L1){hier_label}: исход = {_h_out}, время обучения = {_h_t} мин; "
                f"покрытие Q-таблиц: L3={cov_l3}, L2={cov_l2}, L1={cov_l1}. "
            )
        except Exception:
            hrl_detail = ""
        conclusions.append(
            f"<b>Иерархический RL (3 уровня: L3 НГ → L2 РТП → L1 НБТП).</b> {hrl_detail}"
            f"<u>Преимущества HRL:</u> декомпозиция задачи соответствует реальной вертикали "
            f"управления МЧС; curriculum learning (serp→смешанный→tuapse) ускоряет "
            f"сходимость на сложных сценариях; интринзическая награда L1 стимулирует "
            f"достижение суб-целей независимо от конечного исхода; "
            f"мягкое маскирование действий сокращает пространство поиска. "
            f"<u>Недостатки HRL:</u> увеличенное время обучения (3 уровня × эпизодов), "
            f"необходимость настройки гиперпараметров k2/k3/λ, "
            f"чувствительность к распределению приоров целей L2."
        )
    else:
        conclusions.append(
            "<b>Сравнение Flat RL vs Иерархический RL.</b> "
            "Сравнительный анализ недоступен: необходимо обучить оба агента "
            "(вкладки «🤖 5. Flat RL» и «🏛 6. Иерарх. RL»), затем запустить "
            "массовый эксперимент (вкладка «🔬 7. Массовое»). "
            "<u>Ожидаемые преимущества HRL:</u> лучшая обобщаемость на сложные сценарии, "
            "соответствие вертикали командования МЧС, ускоренная сходимость при обучении. "
            "<u>Ожидаемые преимущества Flat RL:</u> простота и скорость обучения, "
            "прозрачность политики в небольших пространствах состояний."
        )

    # ── Вывод 6 — Нормативное соответствие ────────────────────────────────
    q_req = stats.get("q_foam_required_ls", 0)
    q_act = stats.get("foam_flow_last_ls", 0)
    norm_status = "соответствует" if q_act >= q_req * 0.95 or stats["extinguished"] else "не достигнут"
    conclusions.append(
        f"<b>Нормативный анализ (ГОСТ Р 51043-2002).</b> "
        f"Требуемый расход пенного раствора: {q_req:.0f} л/с; "
        f"фактический расход последней атаки: {q_act:.0f} л/с — норматив <b>{norm_status}</b>. "
        f"Выполнение нормативных требований является необходимым условием учёта результатов "
        f"при планировании реальных противопожарных мероприятий."
    )

    # ── Вывод 7 — Рекомендации и научная новизна ──────────────────────────
    conclusions.append(
        "<b>Научная новизна и перспективы.</b> "
        "Разработанная симуляционная среда с RL-агентом воспроизводит тактику тушения "
        "резервуарных пожаров с учётом нормативной базы МЧС. Применение иерархического "
        "Q-learning с curriculum learning обеспечивает декомпозицию командования "
        "(НГ→РТП→НБТП), что является новым подходом в области автоматизированных систем "
        "управления реагированием при тушении пожаров. "
        "Перспективы: замена табличного Q-learning на Deep Q-Network (DQN/PPO) "
        "для непрерывного пространства состояний; калибровка приоров целей L2 "
        "по базе актов пожаров ВНИИПО; интеграция с ГИС-системами для "
        "ситуационного моделирования в реальном времени."
    )

    for i, txt in enumerate(conclusions, 1):
        # Номер вывода выделен жирным синим
        numbered = f"<b><font color='#1a5276'>{i}.</font></b>  {txt}"
        story.append(_para(numbered, sty["concl"]))
        story.append(Spacer(1, 2*mm))

    story.append(Spacer(1, 6*mm))

    # ═══════════════════════════════════════════════════════════
    # СПИСОК ЛИТЕРАТУРЫ
    # ═══════════════════════════════════════════════════════════
    story.append(_para("Список литературы", sty["h1"]))
    story.append(_hr())
    refs = [
        "[1] ГОСТ Р 51043-2002. Установки водяного и пенного пожаротушения автоматические.",
        "[2] СП 155.13130.2014. Склады нефти и нефтепродуктов. Требования пожарной безопасности.",
        "[3] ГОСТ Р 22.7.01-2021. Безопасность в чрезвычайных ситуациях. Единая система управления.",
        "[4] Watkins C.J.C.H., Dayan P. Q-learning // Machine Learning. 1992. Vol. 8, No. 3–4. P. 279–292.",
        "[5] Sutton R.S., Barto A.G. Reinforcement Learning: An Introduction. MIT Press, 2018.",
        "[6] Dietterich T.G. Hierarchical Reinforcement Learning with the MAXQ Value Function "
            "Decomposition // JAIR. 2000. Vol. 13. P. 227–303.",
        "[7] Bengio Y. et al. Curriculum Learning // ICML. 2009. P. 41–48.",
    ]
    for ref in refs:
        story.append(_para(ref, sty["note"]))

    story.append(Spacer(1, 6*mm))
    story.append(_para(
        f"Отчёт сформирован автоматически системой САУР ПСП v1.0 | {stats['generated_at']}",
        sty["note"]
    ))

    doc.build(story, onFirstPage=page_cb, onLaterPages=page_cb)
    return output_path


# ══════════════════════════════════════════════════════════════════════════════
# ОТЧЁТ РЕЖИМА «ТРЕНАЖЁР РТП»
# ══════════════════════════════════════════════════════════════════════════════

def generate_trainer_report(sim: "TankFireSim",
                             trainer_log: list,
                             trainer_score: int,
                             trainer_steps: int,
                             output_path: str = "") -> str:
    """PDF-отчёт по результатам тренировки в режиме Тренажёр РТП.

    Содержит:
      1. Обложка с итоговым счётом
      2. Итоговая таблица тренировки
      3. Разбор действий по шагам
      4. Динамика пожара (графики)
      5. Итоги пожара
      6. Нормативный анализ
      7. Хронология событий
      8. Рекомендации тренируемому
    """
    stats = collect_stats(sim)

    try:
        from .tank_fire_sim import ACTIONS
    except ImportError:
        from tank_fire_sim import ACTIONS

    if not output_path:
        ts = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        output_path = os.path.join(
            os.path.dirname(__file__),
            f"trainer_report_{stats['scenario_key']}_{ts}.pdf"
        )

    C_TR = colors.HexColor("#8e44ad")      # фиолетовый акцент тренажёра
    C_TR_LIGHT = colors.HexColor("#f3e5f5")

    sty = {
        "title":  ParagraphStyle("tr_title",  fontName=_FONT_B, fontSize=17,
                                  textColor=colors.white, alignment=TA_CENTER,
                                  spaceAfter=6, leading=22),
        "subtitle": ParagraphStyle("tr_sub",  fontName=_FONT_B, fontSize=11,
                                   textColor=C_TR, alignment=TA_CENTER, spaceAfter=4),
        "h1":     ParagraphStyle("tr_h1",     fontName=_FONT_B, fontSize=13,
                                  textColor=C_TR, spaceBefore=12, spaceAfter=5),
        "h2":     ParagraphStyle("tr_h2",     fontName=_FONT_B, fontSize=10,
                                  textColor=colors.HexColor("#2c3e50"),
                                  spaceBefore=8, spaceAfter=3),
        "body":   ParagraphStyle("tr_body",   fontName=_FONT, fontSize=9,
                                  textColor=colors.HexColor("#2c3e50"), leading=14,
                                  spaceAfter=5, alignment=TA_JUSTIFY),
        "note":   ParagraphStyle("tr_note",   fontName=_FONT_I, fontSize=8,
                                  textColor=colors.HexColor("#5d6d7e"), spaceAfter=3),
        "ok":     ParagraphStyle("tr_ok",     fontName=_FONT_B, fontSize=9,
                                  textColor=C_SUCCESS, spaceAfter=3),
        "warn":   ParagraphStyle("tr_warn",   fontName=_FONT_B, fontSize=9,
                                  textColor=C_WARN, spaceAfter=3),
        "score":  ParagraphStyle("tr_score",  fontName=_FONT_B, fontSize=28,
                                  textColor=C_TR, alignment=TA_CENTER),
    }

    # ── Вычисления по журналу тренировки ─────────────────────────────────────
    n_perfect = sum(1 for s in trainer_log if s["pts"] == 20)
    n_good    = sum(1 for s in trainer_log if s["pts"] == 10)
    n_valid   = sum(1 for s in trainer_log if s["pts"] == 2)
    n_wrong   = sum(1 for s in trainer_log if s["pts"] < 0)
    max_score = trainer_steps * 20
    efficiency = int(100 * trainer_score / max_score) if max_score > 0 else 0

    PHASE_NAMES = {"S1":"Обнаружение","S2":"Развёртывание","S3":"Локализация",
                   "S4":"Тушение","S5":"Контроль"}

    def _verdict(pts):
        if pts == 20:  return "✅ Оптимально"
        if pts == 10:  return "👍 Хорошо"
        if pts == 2:   return "⚠ Допустимо"
        return "❌ Ошибка"

    def _page_cb(canv, doc):
        canv.saveState()
        w, h = A4
        canv.setFillColor(C_TR)
        canv.rect(0, h - 16*mm, w, 16*mm, fill=1, stroke=0)
        canv.setFont(_FONT_B, 9)
        canv.setFillColor(colors.white)
        canv.drawString(15*mm, h - 10*mm, "САУР ПСП — Отчёт тренажёра РТП")
        canv.setFont(_FONT, 8)
        canv.drawRightString(w - 15*mm, h - 10*mm, stats["generated_at"])
        canv.setFillColor(C_TR_LIGHT)
        canv.rect(0, 0, w, 11*mm, fill=1, stroke=0)
        canv.setFont(_FONT, 7)
        canv.setFillColor(colors.HexColor("#2c3e50"))
        canv.drawString(15*mm, 3.5*mm, "САУР ПСП v1.0 — Тренажёр руководителя тушения пожара")
        canv.drawRightString(w - 15*mm, 3.5*mm, f"Страница {doc.page}")
        canv.restoreState()

    doc = SimpleDocTemplate(
        output_path, pagesize=A4,
        leftMargin=20*mm, rightMargin=20*mm,
        topMargin=25*mm, bottomMargin=20*mm,
        title=f"Отчёт тренажёра РТП — {stats['scenario_name']}",
        author="САУР ПСП v1.0",
    )
    story: list = []

    # ═══════════════════════════════════════════════════════════
    # 1. ОБЛОЖКА
    # ═══════════════════════════════════════════════════════════
    story.append(Spacer(1, 15*mm))
    story.append(HRFlowable(width="100%", thickness=4, color=C_TR, spaceAfter=4))

    cover_tbl = Table([[Paragraph(
        "🎓  ОТЧЁТ ТРЕНАЖЁРА РТП<br/>"
        "<font size=12>Результаты тренировки руководителя тушения пожара</font>",
        sty["title"]
    )]], colWidths=[170*mm])
    cover_tbl.setStyle(TableStyle([
        ("BACKGROUND",    (0,0), (-1,-1), C_TR),
        ("TOPPADDING",    (0,0), (-1,-1), 16),
        ("BOTTOMPADDING", (0,0), (-1,-1), 16),
        ("LEFTPADDING",   (0,0), (-1,-1), 14),
    ]))
    story.append(cover_tbl)
    story.append(HRFlowable(width="100%", thickness=3,
                             color=colors.HexColor("#e67e22"), spaceAfter=8))

    # Счёт крупно
    story.append(_para(f"Итоговый счёт: {trainer_score} / {max_score}  ({efficiency}%)",
                        sty["score"]))
    story.append(Spacer(1, 4*mm))

    outcome_icon = "✅" if sim.extinguished else ("🔒" if sim.localized else "🔥")
    outcome_sty  = sty["ok"] if sim.extinguished else sty["warn"]
    story.append(_para(
        f"<b>Исход пожара:</b> {outcome_icon} {stats['outcome'].upper()}  |  "
        f"<b>Время:</b> {stats['sim_duration_str']}  |  "
        f"<b>Сценарий:</b> {stats['scenario_name'][:50]}",
        outcome_sty))
    story.append(Spacer(1, 5*mm))

    cover_meta = [
        ["Параметр", "Значение"],
        ["Сценарий",                  stats["scenario_name"][:60]],
        ["Объект",                    stats["rvs_name"]],
        ["Исход пожара",              f"{outcome_icon} {stats['outcome']}"],
        ["Время симуляции",           stats["sim_duration_str"]],
        ["Шагов тренировки",          str(trainer_steps)],
        ["Итоговый счёт",             f"{trainer_score} из {max_score}"],
        ["Эффективность",             f"{efficiency}%"],
        ["Оптимальных действий",      f"{n_perfect} ({int(100*n_perfect/max(trainer_steps,1))}%)"],
        ["Хороших действий (топ-3)",  str(n_good)],
        ["Допустимых",                str(n_valid)],
        ["Ошибочных",                 str(n_wrong)],
        ["Дата",                      stats["generated_at"].replace("T", "  ")],
    ]
    ct = Table(cover_meta, colWidths=[80*mm, 90*mm])
    ct.setStyle(_table_style(1))
    story.append(ct)
    story.append(PageBreak())

    # ═══════════════════════════════════════════════════════════
    # 2. РАЗБОР ДЕЙСТВИЙ ПО ШАГАМ
    # ═══════════════════════════════════════════════════════════
    story.append(_para("1. Пошаговый разбор действий РТП", sty["h1"]))
    story.append(_hr())
    story.append(_para(
        "В таблице приведён разбор каждого шага тренировки: "
        "действие, выбранное обучаемым, оптимальное действие агента, "
        "Q-ценности и оценка.",
        sty["body"]
    ))

    hdr = [["Ч+", "Фаза", "Действие РТП", "Оптим. действие",
            "Q (РТП)", "Q (опт.)", "Баллы", "Оценка"]]
    rows = []
    for s in trainer_log:
        ua = s["user_a"];  ba = s["best_a"]
        rows.append([
            str(s["t"]),
            PHASE_NAMES.get(s["phase"], s["phase"]),
            f"[{ACTIONS[ua][0]}] {ACTIONS[ua][2][:28]}",
            f"[{ACTIONS[ba][0]}] {ACTIONS[ba][2][:28]}",
            f"{s['q_user']:+.2f}",
            f"{s['q_best']:+.2f}",
            f"{s['pts']:+d}",
            _verdict(s["pts"]),
        ])

    step_tbl = Table(hdr + rows,
                     colWidths=[12*mm, 20*mm, 42*mm, 42*mm,
                                15*mm, 15*mm, 12*mm, 22*mm])
    ts = _table_style(1)
    # Раскрасить строки по оценке
    for i, s in enumerate(trainer_log, start=1):
        if s["pts"] == 20:
            bg = colors.HexColor("#d5f5e3")
        elif s["pts"] == 10:
            bg = colors.HexColor("#fef9e7")
        elif s["pts"] < 0:
            bg = colors.HexColor("#fde8e8")
        else:
            bg = colors.white
        ts.add("BACKGROUND", (0, i), (-1, i), bg)
    step_tbl.setStyle(ts)
    story.append(step_tbl)
    story.append(Spacer(1, 5*mm))

    # ═══════════════════════════════════════════════════════════
    # 3. ДИНАМИКА ПОЖАРА
    # ═══════════════════════════════════════════════════════════
    story.append(_para("2. Динамика показателей пожара", sty["h1"]))
    story.append(_hr())
    if sim.h_fire:
        img_buf = _build_metrics_figure(sim)
        story.append(RLImage(img_buf, width=170*mm, height=90*mm))
    story.append(Spacer(1, 3*mm))

    # ═══════════════════════════════════════════════════════════
    # 4. ИТОГИ ПОЖАРА
    # ═══════════════════════════════════════════════════════════
    story.append(_para("3. Итоги пожара", sty["h1"]))
    story.append(_hr())
    outcome_data = [
        ["Показатель", "Значение"],
        ["Исход",                   f"{outcome_icon} {stats['outcome'].upper()}"],
        ["Время до завершения",     stats["sim_duration_str"]],
        ["Нач. площадь пожара",     f"{stats['fire_area_initial_m2']:.0f} м²"],
        ["Конечная площадь",        f"{stats['fire_area_final_m2']:.0f} м²"],
        ["Макс. площадь",           f"{stats['fire_area_max_m2']:.0f} м²"],
        ["Пенных атак",             str(stats["foam_attacks_total"])],
        ["Расход ОВ (макс.)",       f"{stats['water_flow_max_ls']:.0f} л/с"],
        ["Макс. уровень риска",     f"{stats['risk_max']:.2f}"],
        ["Стволов на РВС (макс.)",  str(stats["trunks_burn_max"])],
    ]
    ot = Table(outcome_data, colWidths=[85*mm, 85*mm])
    ot.setStyle(_table_style(1))
    story.append(ot)
    story.append(Spacer(1, 5*mm))

    # ═══════════════════════════════════════════════════════════
    # 5. НОРМАТИВНЫЙ АНАЛИЗ
    # ═══════════════════════════════════════════════════════════
    story.append(_para("4. Нормативный анализ (ГОСТ Р 51043-2002)", sty["h1"]))
    story.append(_hr())
    q_fr = stats["q_foam_required_ls"];  q_fa = stats.get("foam_flow_last_ls", 0)
    q_cr = stats["q_cooling_required_ls"]; q_ca = stats.get("water_flow_mean_ls", 0)
    roof  = stats.get("roof_obstruction", 0)
    def _ns(a, r): return "✅ Выполнено" if r > 0 and a >= r*0.95 else "❌ Не выполнено"
    norm_data = [
        ["Требование", "Требуемое", "Фактическое", "Статус"],
        ["Расход пенного р-ра",      f"{q_fr:.1f} л/с", f"{q_fa:.1f} л/с", _ns(q_fa, q_fr)],
        ["Расход охлаждения РВС",    f"{q_cr:.1f} л/с", f"{q_ca:.1f} л/с", _ns(q_ca, q_cr)],
        ["Препятствие крыши",        "≤ 30%",            f"{roof*100:.0f}%",
         "✅ Норма" if roof <= 0.30 else "⚠ Превышение"],
        ["Число пенных атак",        "≥ 1",
         str(stats["foam_attacks_total"]),
         "✅ Выполнено" if stats["foam_attacks_total"] >= 1 else "❌ Нет"],
    ]
    nt = Table(norm_data, colWidths=[75*mm, 33*mm, 33*mm, 29*mm])
    nt.setStyle(_table_style(1))
    story.append(nt)
    story.append(Spacer(1, 5*mm))

    # ═══════════════════════════════════════════════════════════
    # 6. ХРОНОЛОГИЯ СОБЫТИЙ
    # ═══════════════════════════════════════════════════════════
    story.append(_para("5. Хронология событий", sty["h1"]))
    story.append(_hr())
    try:
        from .tank_fire_sim import P as _P
    except ImportError:
        from tank_fire_sim import P as _P
    _CLR = {_P["danger"]:"❗ ", _P["warn"]:"⚠ ", _P["success"]:"✅ ", _P["info"]:"ℹ "}
    ev_rows = [["Время, мин", "Событие"]]
    for ev in stats["events_log"][:50]:
        ev_rows.append([str(ev["t_min"]),
                        (_CLR.get(ev["color"], "  ") + ev["text"])[:100]])
    et = Table(ev_rows, colWidths=[25*mm, 145*mm])
    et.setStyle(_table_style(1))
    story.append(et)
    story.append(Spacer(1, 5*mm))

    # ═══════════════════════════════════════════════════════════
    # 7. РЕКОМЕНДАЦИИ ТРЕНИРУЕМОМУ
    # ═══════════════════════════════════════════════════════════
    story.append(_para("6. Рекомендации по итогам тренировки", sty["h1"]))
    story.append(_hr())

    recom = []
    if efficiency >= 80:
        recom.append(("✅ Высокий уровень подготовки.",
                      f"Эффективность {efficiency}%. Большинство решений совпадают с "
                      "оптимальной стратегией агента. Рекомендуется усложнить сценарий "
                      "(ранг №4, плавающая крыша, высокий начальный уровень риска)."))
    elif efficiency >= 50:
        recom.append(("⚠ Средний уровень подготовки.",
                      f"Эффективность {efficiency}%. Часть решений субоптимальна. "
                      "Рекомендуется повторить сценарий с акцентом на фазы, "
                      "в которых допущены ошибки."))
    else:
        recom.append(("❌ Требуется дополнительная подготовка.",
                      f"Эффективность {efficiency}%. Значительная часть решений "
                      "расходится с оптимальной стратегией. "
                      "Рекомендуется изучить Справочник действий по фазам."))

    if n_wrong > 0:
        wrong_actions = [ACTIONS[s["user_a"]][0]
                         for s in trainer_log if s["pts"] < 0]
        recom.append(("📋 Ошибочные действия:",
                      "Действия, недопустимые в текущей фазе: "
                      + ", ".join(f"[{a}]" for a in wrong_actions)
                      + ". Изучите таблицу допустимых действий по фазам."))

    missed_phases = set(s["phase"] for s in trainer_log
                        if s["user_a"] != s["best_a"])
    if missed_phases:
        recom.append(("📌 Фазы, требующие доработки:",
                      "В следующих фазах выбор расходился с оптимальным: "
                      + ", ".join(f"{p}" for p in sorted(missed_phases))
                      + ". Повторите раздел «Справочник действий» для этих фаз."))

    for title, body in recom:
        story.append(_para(f"<b>{title}</b>  {body}", sty["body"]))

    story.append(Spacer(1, 4*mm))
    story.append(_para(
        f"Отчёт сформирован автоматически системой САУР ПСП v1.0 | {stats['generated_at']}",
        sty["note"]
    ))

    doc.build(story, onFirstPage=_page_cb, onLaterPages=_page_cb)
    return output_path


# ══════════════════════════════════════════════════════════════════════════════
# ОТЧЁТ РЕЖИМА «СППР»
# ══════════════════════════════════════════════════════════════════════════════

def generate_sppр_report(sim: "TankFireSim",
                          sppр_log: list,
                          sppр_total: int,
                          sppр_deviations: int,
                          output_path: str = "") -> str:
    """PDF-отчёт по результатам сессии в режиме СППР.

    Содержит:
      1. Обложка с показателями сессии
      2. Итоговая статистика (принятые/отклонённые рекомендации)
      3. Журнал решений оперативного штаба
      4. Динамика пожара (графики)
      5. Итоги пожара
      6. Нормативный анализ
      7. Хронология событий
      8. Выводы о взаимодействии «человек — СППР»
    """
    stats = collect_stats(sim)

    try:
        from .tank_fire_sim import ACTIONS
    except ImportError:
        from tank_fire_sim import ACTIONS

    if not output_path:
        ts = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        output_path = os.path.join(
            os.path.dirname(__file__),
            f"sppр_report_{stats['scenario_key']}_{ts}.pdf"
        )

    C_SP = colors.HexColor("#1a6da8")      # синий акцент СППР
    C_SP_LIGHT = colors.HexColor("#e3f0fb")

    sty = {
        "title":   ParagraphStyle("sp_title",  fontName=_FONT_B, fontSize=17,
                                   textColor=colors.white, alignment=TA_CENTER,
                                   spaceAfter=6, leading=22),
        "subtitle":ParagraphStyle("sp_sub",    fontName=_FONT_B, fontSize=11,
                                   textColor=C_SP, alignment=TA_CENTER, spaceAfter=4),
        "h1":      ParagraphStyle("sp_h1",     fontName=_FONT_B, fontSize=13,
                                   textColor=C_SP, spaceBefore=12, spaceAfter=5),
        "h2":      ParagraphStyle("sp_h2",     fontName=_FONT_B, fontSize=10,
                                   textColor=colors.HexColor("#2c3e50"),
                                   spaceBefore=8, spaceAfter=3),
        "body":    ParagraphStyle("sp_body",   fontName=_FONT, fontSize=9,
                                   textColor=colors.HexColor("#2c3e50"), leading=14,
                                   spaceAfter=5, alignment=TA_JUSTIFY),
        "note":    ParagraphStyle("sp_note",   fontName=_FONT_I, fontSize=8,
                                   textColor=colors.HexColor("#5d6d7e"), spaceAfter=3),
        "ok":      ParagraphStyle("sp_ok",     fontName=_FONT_B, fontSize=9,
                                   textColor=C_SUCCESS, spaceAfter=3),
        "warn":    ParagraphStyle("sp_warn",   fontName=_FONT_B, fontSize=9,
                                   textColor=C_WARN, spaceAfter=3),
        "pct":     ParagraphStyle("sp_pct",    fontName=_FONT_B, fontSize=28,
                                   textColor=C_SP, alignment=TA_CENTER),
    }

    accepted  = sppр_total - sppр_deviations
    agree_pct = int(100 * accepted / max(sppр_total, 1))
    PHASE_NAMES = {"S1":"Обнаружение","S2":"Развёртывание","S3":"Локализация",
                   "S4":"Тушение","S5":"Контроль"}

    def _page_cb(canv, doc):
        canv.saveState()
        w, h = A4
        canv.setFillColor(C_SP)
        canv.rect(0, h - 16*mm, w, 16*mm, fill=1, stroke=0)
        canv.setFont(_FONT_B, 9)
        canv.setFillColor(colors.white)
        canv.drawString(15*mm, h - 10*mm, "САУР ПСП — Отчёт режима СППР")
        canv.setFont(_FONT, 8)
        canv.drawRightString(w - 15*mm, h - 10*mm, stats["generated_at"])
        canv.setFillColor(C_SP_LIGHT)
        canv.rect(0, 0, w, 11*mm, fill=1, stroke=0)
        canv.setFont(_FONT, 7)
        canv.setFillColor(colors.HexColor("#2c3e50"))
        canv.drawString(15*mm, 3.5*mm,
                        "САУР ПСП v1.0 — Система поддержки принятия решений оперативного штаба")
        canv.drawRightString(w - 15*mm, 3.5*mm, f"Страница {doc.page}")
        canv.restoreState()

    doc = SimpleDocTemplate(
        output_path, pagesize=A4,
        leftMargin=20*mm, rightMargin=20*mm,
        topMargin=25*mm, bottomMargin=20*mm,
        title=f"Отчёт СППР — {stats['scenario_name']}",
        author="САУР ПСП v1.0",
    )
    story: list = []

    # ═══════════════════════════════════════════════════════════
    # 1. ОБЛОЖКА
    # ═══════════════════════════════════════════════════════════
    story.append(Spacer(1, 15*mm))
    story.append(HRFlowable(width="100%", thickness=4, color=C_SP, spaceAfter=4))

    cover_tbl = Table([[Paragraph(
        "🧭  ОТЧЁТ СЕССИИ СППР<br/>"
        "<font size=12>Анализ решений оперативного штаба "
        "при поддержке системы рекомендаций</font>",
        sty["title"]
    )]], colWidths=[170*mm])
    cover_tbl.setStyle(TableStyle([
        ("BACKGROUND",    (0,0), (-1,-1), C_SP),
        ("TOPPADDING",    (0,0), (-1,-1), 16),
        ("BOTTOMPADDING", (0,0), (-1,-1), 16),
        ("LEFTPADDING",   (0,0), (-1,-1), 14),
    ]))
    story.append(cover_tbl)
    story.append(HRFlowable(width="100%", thickness=3,
                             color=colors.HexColor("#e67e22"), spaceAfter=8))

    story.append(_para(f"Согласие с рекомендациями СППР: {agree_pct}%", sty["pct"]))
    story.append(Spacer(1, 4*mm))

    outcome_icon = "✅" if sim.extinguished else ("🔒" if sim.localized else "🔥")
    outcome_sty  = sty["ok"] if sim.extinguished else sty["warn"]
    story.append(_para(
        f"<b>Исход пожара:</b> {outcome_icon} {stats['outcome'].upper()}  |  "
        f"<b>Время:</b> {stats['sim_duration_str']}  |  "
        f"<b>Сценарий:</b> {stats['scenario_name'][:50]}",
        outcome_sty))
    story.append(Spacer(1, 5*mm))

    cover_meta = [
        ["Параметр", "Значение"],
        ["Сценарий",                  stats["scenario_name"][:60]],
        ["Объект",                    stats["rvs_name"]],
        ["Исход пожара",              f"{outcome_icon} {stats['outcome']}"],
        ["Время симуляции",           stats["sim_duration_str"]],
        ["Всего шагов управления",    str(sppр_total)],
        ["Принято рекомендаций",      f"{accepted} ({agree_pct}%)"],
        ["Отклонено (переопределено)",f"{sppр_deviations} ({100-agree_pct}%)"],
        ["Дата сессии",               stats["generated_at"].replace("T", "  ")],
    ]
    ct = Table(cover_meta, colWidths=[85*mm, 85*mm])
    ct.setStyle(_table_style(1))
    story.append(ct)
    story.append(PageBreak())

    # ═══════════════════════════════════════════════════════════
    # 2. ЖУРНАЛ РЕШЕНИЙ ОПЕРАТИВНОГО ШТАБА
    # ═══════════════════════════════════════════════════════════
    story.append(_para("1. Журнал решений оперативного штаба", sty["h1"]))
    story.append(_hr())
    story.append(_para(
        "В таблице отражены все шаги управления: "
        "рекомендация системы СППР, фактическое решение штаба и статус принятия.",
        sty["body"]
    ))

    # Из _sppр_log берём только записи с override (accepted=False)
    overrides = [e for e in sppр_log if not e.get("accepted", True)]

    if overrides:
        hdr = [["Ч+", "Фаза", "Рекомендация СППР", "Решение штаба", "Статус"]]
        rows = []
        for e in overrides:
            rec = e.get("rec_a", 0);  usr = e.get("user_a", 0)
            rows.append([
                str(e.get("t", "—")),
                PHASE_NAMES.get(e.get("phase",""), e.get("phase","")),
                f"[{ACTIONS[rec][0]}] {ACTIONS[rec][2][:35]}",
                f"[{ACTIONS[usr][0]}] {ACTIONS[usr][2][:35]}",
                "⚡ Переопределено",
            ])
        dec_tbl = Table(hdr + rows,
                        colWidths=[12*mm, 22*mm, 52*mm, 52*mm, 32*mm])
        ts = _table_style(1)
        for i in range(1, len(rows)+1):
            ts.add("BACKGROUND", (0,i), (-1,i), colors.HexColor("#fff3cd"))
        dec_tbl.setStyle(ts)
        story.append(dec_tbl)
    else:
        story.append(_para(
            "✅ Все рекомендации системы СППР приняты штабом без переопределений.",
            sty["ok"]
        ))
    story.append(Spacer(1, 3*mm))
    story.append(_para(
        f"Принято автоматически: {accepted} из {sppр_total} шагов ({agree_pct}%). "
        f"Переопределено штабом: {sppр_deviations}.",
        sty["note"]
    ))
    story.append(Spacer(1, 5*mm))

    # ═══════════════════════════════════════════════════════════
    # 3. ДИНАМИКА ПОЖАРА
    # ═══════════════════════════════════════════════════════════
    story.append(_para("2. Динамика показателей пожара", sty["h1"]))
    story.append(_hr())
    if sim.h_fire:
        img_buf = _build_metrics_figure(sim)
        story.append(RLImage(img_buf, width=170*mm, height=90*mm))
    story.append(Spacer(1, 3*mm))

    # ═══════════════════════════════════════════════════════════
    # 4. ИТОГИ ПОЖАРА
    # ═══════════════════════════════════════════════════════════
    story.append(_para("3. Итоги пожара", sty["h1"]))
    story.append(_hr())
    outcome_data = [
        ["Показатель", "Значение"],
        ["Исход",                   f"{outcome_icon} {stats['outcome'].upper()}"],
        ["Время до завершения",     stats["sim_duration_str"]],
        ["Нач. площадь пожара",     f"{stats['fire_area_initial_m2']:.0f} м²"],
        ["Конечная площадь",        f"{stats['fire_area_final_m2']:.0f} м²"],
        ["Пенных атак",             str(stats["foam_attacks_total"])],
        ["Расход ОВ (макс.)",       f"{stats['water_flow_max_ls']:.0f} л/с"],
        ["Макс. уровень риска",     f"{stats['risk_max']:.2f}"],
        ["Боевых участков",         str(stats["n_bu"])],
    ]
    ot = Table(outcome_data, colWidths=[85*mm, 85*mm])
    ot.setStyle(_table_style(1))
    story.append(ot)
    story.append(Spacer(1, 5*mm))

    # ═══════════════════════════════════════════════════════════
    # 5. НОРМАТИВНЫЙ АНАЛИЗ
    # ═══════════════════════════════════════════════════════════
    story.append(_para("4. Нормативный анализ (ГОСТ Р 51043-2002)", sty["h1"]))
    story.append(_hr())
    q_fr = stats["q_foam_required_ls"];  q_fa = stats.get("foam_flow_last_ls", 0)
    q_cr = stats["q_cooling_required_ls"]; q_ca = stats.get("water_flow_mean_ls", 0)
    roof  = stats.get("roof_obstruction", 0)
    def _ns2(a, r): return "✅ Выполнено" if r > 0 and a >= r*0.95 else "❌ Не выполнено"
    norm_data = [
        ["Требование", "Требуемое", "Фактическое", "Статус"],
        ["Расход пенного р-ра",    f"{q_fr:.1f} л/с", f"{q_fa:.1f} л/с", _ns2(q_fa, q_fr)],
        ["Расход охлаждения РВС",  f"{q_cr:.1f} л/с", f"{q_ca:.1f} л/с", _ns2(q_ca, q_cr)],
        ["Препятствие крыши",      "≤ 30%",            f"{roof*100:.0f}%",
         "✅ Норма" if roof <= 0.30 else "⚠ Превышение"],
        ["Число пенных атак",      "≥ 1",
         str(stats["foam_attacks_total"]),
         "✅ Выполнено" if stats["foam_attacks_total"] >= 1 else "❌ Нет"],
    ]
    nt = Table(norm_data, colWidths=[75*mm, 33*mm, 33*mm, 29*mm])
    nt.setStyle(_table_style(1))
    story.append(nt)
    story.append(Spacer(1, 5*mm))

    # ═══════════════════════════════════════════════════════════
    # 6. ХРОНОЛОГИЯ СОБЫТИЙ
    # ═══════════════════════════════════════════════════════════
    story.append(_para("5. Хронология событий", sty["h1"]))
    story.append(_hr())
    try:
        from .tank_fire_sim import P as _P2
    except ImportError:
        from tank_fire_sim import P as _P2
    _CLR2 = {_P2["danger"]:"❗ ", _P2["warn"]:"⚠ ", _P2["success"]:"✅ ", _P2["info"]:"ℹ "}
    ev_rows = [["Время, мин", "Событие"]]
    for ev in stats["events_log"][:50]:
        ev_rows.append([str(ev["t_min"]),
                        (_CLR2.get(ev["color"], "  ") + ev["text"])[:100]])
    et = Table(ev_rows, colWidths=[25*mm, 145*mm])
    et.setStyle(_table_style(1))
    story.append(et)
    story.append(Spacer(1, 5*mm))

    # ═══════════════════════════════════════════════════════════
    # 7. ВЫВОДЫ: ВЗАИМОДЕЙСТВИЕ «ЧЕЛОВЕК — СППР»
    # ═══════════════════════════════════════════════════════════
    story.append(_para("6. Выводы о взаимодействии «Штаб — СППР»", sty["h1"]))
    story.append(_hr())

    concl = []
    # Вывод 1: исход пожара
    if sim.extinguished:
        concl.append(
            f"<b>Исход.</b> Пожар успешно ликвидирован за {stats['sim_duration_str']}. "
            "Совместная работа оперативного штаба и системы СППР обеспечила "
            "достижение цели управления.")
    elif sim.localized:
        concl.append(
            f"<b>Исход.</b> Пожар локализован за {stats['sim_duration_str']}. "
            "Дальнейшие действия по ликвидации горения не завершены.")
    else:
        concl.append(
            f"<b>Исход.</b> Пожар не ликвидирован в отведённое время "
            f"({stats['sim_duration_str']}). Рекомендуется пересмотреть "
            "стратегию управления ресурсами.")

    # Вывод 2: согласие с СППР
    if agree_pct >= 80:
        concl.append(
            f"<b>Взаимодействие со СППР.</b> Высокий уровень согласия "
            f"({agree_pct}% принятых рекомендаций). Штаб в основном следовал "
            "рекомендациям системы, что свидетельствует о доверии к СППР и "
            "качестве рекомендаций агента.")
    elif agree_pct >= 50:
        concl.append(
            f"<b>Взаимодействие со СППР.</b> Умеренное согласие "
            f"({agree_pct}%). Штаб переопределял рекомендации в "
            f"{sppр_deviations} из {sppр_total} случаев. Рекомендуется "
            "проанализировать причины отклонений.")
    else:
        concl.append(
            f"<b>Взаимодействие со СППР.</b> Низкий уровень согласия "
            f"({agree_pct}%). Большинство рекомендаций отклонено. "
            "Возможные причины: недоверие к модели, специфика сценария, "
            "или необходимость доработки функции награды агента.")

    # Вывод 3: нормативное соответствие
    norm_ok = (q_fa >= q_fr * 0.95 or sim.extinguished)
    concl.append(
        f"<b>Нормативное соответствие.</b> "
        + ("Требования ГОСТ Р 51043-2002 по расходу пенного раствора "
           f"(Q ≥ {q_fr:.0f} л/с) выполнены." if norm_ok
           else f"Требуемый расход пенного раствора ({q_fr:.0f} л/с) "
                f"не достигнут (факт: {q_fa:.0f} л/с). "
                "Необходимо увеличить число ПНС и стволов."))

    for i, txt in enumerate(concl, 1):
        story.append(_para(f"{i}. {txt}", sty["body"]))

    story.append(Spacer(1, 4*mm))
    story.append(_para(
        f"Отчёт сформирован автоматически системой САУР ПСП v1.0 | {stats['generated_at']}",
        sty["note"]
    ))

    doc.build(story, onFirstPage=_page_cb, onLaterPages=_page_cb)
    return output_path


# ══════════════════════════════════════════════════════════════════════════════
# ВЫГРУЗКА ДЛЯ НАУЧНОЙ СТАТЬИ — JSON
# ══════════════════════════════════════════════════════════════════════════════

def export_for_article_json(sim: "TankFireSim", output_path: str = "") -> str:
    """Экспортировать структурированные данные для написания научной статьи в JSON.

    Структура JSON следует шаблону описания результатов вычислительного эксперимента:
      - metadata: описание метода, ссылки, параметры
      - scenario: параметры сценария
      - results: итоговые показатели
      - time_series: временны́е ряды
      - rl_statistics: статистика RL-агента
      - normative_analysis: анализ соответствия нормам
    """
    stats = collect_stats(sim)

    if not output_path:
        ts = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        output_path = os.path.join(
            os.path.dirname(__file__),
            f"article_data_{stats['scenario_key']}_{ts}.json"
        )

    article_data = {
        "metadata": {
            "title": "Результаты вычислительного эксперимента — САУР ПСП",
            "description": (
                "Данные моделирования для разработки раздела «Результаты» "
                "в научной статье по теме: «Адаптивная система управления ресурсами "
                "пожарно-спасательного подразделения на пожарах нефтепродуктов»"
            ),
            "method": "Дискретно-событийное моделирование (DES) + Q-learning (RL)",
            "software": "САУР ПСП v1.0 (Python 3.12, NumPy, ReportLab)",
            "normative_base": [
                "ГОСТ Р 51043-2002 (Изменение №1, 2020)",
                "СП 155.13130.2014",
                "Справочник РТП (ВНИИПО, 2021)",
            ],
            "generated_at": stats["generated_at"],
            "cite_as": (
                "САУР ПСП v1.0 // GitHub: github.com/drivescienceai/saur_sim"
            ),
        },
        "scenario": {
            "name":          stats["scenario_name"],
            "key":           stats["scenario_key"],
            "object":        stats["rvs_name"],
            "fuel":          stats["fuel"],
            "fire_rank":     stats["fire_rank"],
            "fire_area_m2":  stats["fire_area_initial_m2"],
            "foam_intensity_ls_m2_s": stats["foam_intensity_norm"],
            "q_foam_required_ls":     stats["q_foam_required_ls"],
            "q_cooling_required_ls":  stats["q_cooling_required_ls"],
            "roof_obstruction_initial": sim._cfg["roof_obstruction_init"],
        },
        "results": {
            "outcome":           stats["outcome"],
            "extinguished":      stats["extinguished"],
            "localized":         stats["localized"],
            "duration_min":      stats["sim_duration_min"],
            "duration_str":      stats["sim_duration_str"],
            "final_phase":       stats["final_phase"],
            "foam_attacks":      stats["foam_attacks_total"],
            "akp50_used":        stats["akp50_used"],
            "roof_obstruction_final": stats["roof_obstruction"],
            "foam_flow_final_ls":     stats["foam_flow_last_ls"],
            "fire_area_max_m2":  stats["fire_area_max_m2"],
            "fire_area_final_m2":stats["fire_area_final_m2"],
            "water_flow_max_ls": stats["water_flow_max_ls"],
            "water_flow_mean_ls":stats["water_flow_mean_ls"],
            "trunks_max":        stats["trunks_burn_max"],
            "n_pns":             stats["n_pns"],
            "n_bu":              stats["n_bu"],
            "secondary_fire":    stats["secondary_fire"],
            "spill":             stats["spill"],
            "risk_max":          round(stats["risk_max"], 3),
            "risk_mean":         round(stats["risk_mean"], 3),
            "risk_final":        round(stats["risk_final"], 3),
        },
        "time_series": {
            "description": "Временны́е ряды ключевых показателей (шаг 5 мин)",
            "t_min":         [x[0] for x in sim.h_fire],
            "fire_area_m2":  [round(x[1], 1) for x in sim.h_fire],
            "water_flow_ls": [round(x[1], 1) for x in sim.h_water],
            "n_trunks":      [x[1] for x in sim.h_trunks],
            "risk_index":    [round(x[1], 3) for x in sim.h_risk],
            "reward":        [round(r, 3) for r in sim.h_reward],
        },
        "rl_statistics": {
            "description": "Статистика Q-learning агента",
            "algorithm":   "Q-learning (tabular), ε-greedy",
            "state_space": 128,
            "action_space": 15,
            "alpha":       sim.agent.alpha,
            "gamma":       sim.agent.gamma,
            "epsilon_final": round(sim.agent.epsilon, 3),
            "total_steps": stats["rl_steps"],
            "total_episodes": stats["rl_episodes"],
            "total_reward":  round(stats["rl_total_reward"], 2),
            "episode_rewards": [round(r, 2) for r in sim.agent.episode_rewards[-50:]],
            "action_distribution": {
                code: {
                    "count": info["count"],
                    "fraction": round(info["fraction"], 3),
                    "description": info["description"],
                }
                for code, info in stats["action_distribution"].items()
            },
            "q_table_nonzero_cells": int(np.count_nonzero(sim.agent.Q)),
            "q_table_max":  float(sim.agent.Q.max()),
            "q_table_mean": float(sim.agent.Q.mean()),
        },
        "normative_analysis": {
            "description": (
                "Анализ соответствия ГОСТ Р 51043-2002 и СП 155.13130.2014"
            ),
            "foam_attack_physics": {
                "model": "Q_эфф = Q_total × (1 - roof_obstruction)",
                "criterion": "Q_эфф ≥ I_норм × S_пожара",
                "source": "ГОСТ Р 51043-2002, п. 5.1.1.3",
            },
            "cooling_norm_ls_m": 0.8,
            "foam_norm_ls_m2_s": stats["foam_intensity_norm"],
            "foam_reserve_coeff": 5,
            "foam_application_time_min": 15,
        },
    }

    with open(output_path, "w", encoding="utf-8") as f:
        json.dump(article_data, f, ensure_ascii=False, indent=2)

    return output_path


# ══════════════════════════════════════════════════════════════════════════════
# ВЫГРУЗКА ДЛЯ НАУЧНОЙ СТАТЬИ — DOCX
# ══════════════════════════════════════════════════════════════════════════════

def export_for_article_docx(sim: "TankFireSim", output_path: str = "") -> str:
    """Экспортировать резюме результатов в DOCX-формате для написания статьи.

    Документ содержит готовые фрагменты текста для разделов статьи:
    «Объект и метод», «Результаты», «Обсуждение», таблицы и подписи рисунков.
    """
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    stats = collect_stats(sim)

    if not output_path:
        ts = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        output_path = os.path.join(
            os.path.dirname(__file__),
            f"article_draft_{stats['scenario_key']}_{ts}.docx"
        )

    doc = Document()

    # Настройка полей
    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(1.5)

    def _heading(text, level=1):
        p = doc.add_heading(text, level=level)
        p.runs[0].font.color.rgb = RGBColor(0x1a, 0x1f, 0x2e)

    def _para_doc(text, bold=False, italic=False, size=10):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold   = bold
        run.italic = italic
        run.font.size = Pt(size)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        return p

    def _table_doc(headers, rows, col_widths=None):
        tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
        tbl.style = "Table Grid"
        hdr_cells = tbl.rows[0].cells
        for i, h in enumerate(headers):
            hdr_cells[i].text = h
            hdr_cells[i].paragraphs[0].runs[0].bold = True
            hdr_cells[i].paragraphs[0].runs[0].font.size = Pt(9)
        for row_data in rows:
            row_cells = tbl.add_row().cells
            for i, val in enumerate(row_data):
                row_cells[i].text = str(val)
                row_cells[i].paragraphs[0].runs[0].font.size = Pt(9)
        return tbl

    # Заголовок документа
    doc.add_heading("РЕЗУЛЬТАТЫ ВЫЧИСЛИТЕЛЬНОГО ЭКСПЕРИМЕНТА", 0)
    _para_doc(f"Система САУР ПСП v1.0 | Сценарий: {stats['scenario_name']}")
    _para_doc(f"Дата: {stats['generated_at'].replace('T', ' ')}", italic=True, size=9)
    doc.add_paragraph()

    # ── Раздел: Объект и методология
    _heading("1. Объект исследования и методология")
    _para_doc(
        f"Объектом моделирования являлся резервуарный парк с вертикальными "
        f"стальными резервуарами (РВС). В рамках сценария «{stats['scenario_name']}» "
        f"рассматривался пожар {stats['rvs_name']} с горючим «{stats['fuel']}», "
        f"ранг пожара №{stats['fire_rank']}, начальная площадь зеркала горения "
        f"S₀ = {stats['fire_area_initial_m2']:.0f} м²."
    )
    _para_doc(
        "Для моделирования применена система дискретно-событийного моделирования (DES) "
        "с Q-learning агентом управления реагированием. Физическая модель пенного тушения "
        "реализована на основе ГОСТ Р 51043-2002 и СП 155.13130.2014."
    )
    doc.add_paragraph()

    # Таблица нормативных параметров
    _heading("1.1. Нормативные требования", level=2)
    _table_doc(
        ["Параметр", "Значение", "Источник"],
        [
            ["Нормативная интенсивность подачи пены",
             f"{stats['foam_intensity_norm']:.3f} л/(м²·с)",
             "СП 155.13130.2014, табл. 8"],
            ["Требуемый расход пенного раствора",
             f"Q_пена = {stats['q_foam_required_ls']:.1f} л/с",
             "Q = I × S"],
            ["Требуемый расход охлаждения",
             f"Q_охл = {stats['q_cooling_required_ls']:.1f} л/с",
             "q = 0.8 л/(с·м)"],
            ["Нормативное время подачи пены",
             "15 мин",
             "Справочник РТП, стр. 104"],
            ["Коэффициент запаса ОВ",
             "Кз = 5",
             "Справочник РТП, стр. 106"],
        ]
    )
    doc.add_paragraph()

    # ── Раздел: Результаты
    _heading("2. Результаты моделирования")
    _para_doc(
        f"Симуляция продолжалась {stats['sim_duration_str']} ({stats['sim_duration_min']} мин). "
        f"Итог: пожар {stats['outcome']}. "
        f"Финальная фаза: {stats['final_phase']}. "
        f"Проведено пенных атак: {stats['foam_attacks_total']}."
    )
    doc.add_paragraph()

    _heading("2.1. Ключевые показатели", level=2)
    _table_doc(
        ["Показатель", "Значение"],
        [
            ["Исход", stats["outcome"].title()],
            ["Продолжительность", stats["sim_duration_str"]],
            ["Площадь пожара (нач.)", f"{stats['fire_area_initial_m2']:.0f} м²"],
            ["Площадь пожара (макс.)", f"{stats['fire_area_max_m2']:.0f} м²"],
            ["Площадь пожара (фин.)", f"{stats['fire_area_final_m2']:.0f} м²"],
            ["Число пенных атак", str(stats["foam_attacks_total"])],
            ["АКП-50 задействован", "Да" if stats["akp50_used"] else "Нет"],
            ["Препятствие крыши (фин.)", f"{stats['roof_obstruction']*100:.0f}%"],
            ["Расход ОВ (макс.)", f"{stats['water_flow_max_ls']:.0f} л/с"],
            ["Число стволов охлаждения (макс.)", str(stats["trunks_burn_max"])],
            ["Индекс риска (макс.)", f"{stats['risk_max']:.3f}"],
            ["Индекс риска (средн.)", f"{stats['risk_mean']:.3f}"],
        ]
    )
    doc.add_paragraph()

    # ── Раздел: RL-агент
    _heading("3. Результаты работы RL-агента")
    _para_doc(
        f"Q-learning агент с ε-greedy стратегией (α={sim.agent.alpha:.2f}, "
        f"γ={sim.agent.gamma:.2f}) за {stats['rl_steps']} шагов "
        f"({stats['rl_episodes']} эпизодов) набрал суммарную награду "
        f"{stats['rl_total_reward']:.1f}. Финальный ε = {stats['rl_epsilon']:.3f}."
    )
    _para_doc(
        "Наиболее часто выбираемые действия (топ-5):"
    )

    top5 = list(stats["action_distribution"].items())[:5]
    _table_doc(
        ["Код", "Описание", "Кол-во", "Доля, %"],
        [[c, info["description"][:50], info["count"],
          f"{info['fraction']*100:.1f}"]
         for c, info in top5]
    )
    doc.add_paragraph()

    # ── Раздел: Обсуждение
    _heading("4. Обсуждение результатов")
    disc = []
    disc.append(
        "Физическая модель тушения, реализованная на основе формулы "
        "Q_эфф = Q_total × (1 – k_обструкции), позволяет корректно воспроизводить "
        "эффект блокировки пены каркасом плавающей крыши — явление, зафиксированное "
        "на реальных пожарах РВС с плавающей крышей (многократные неудачные пенные атаки)."
    )
    if stats["akp50_used"]:
        disc.append(
            "Применение АКП-50 для подачи ГПС-1000 через верхний люк снижает "
            "долю обструкции с 70% до 20%, что обеспечивает достаточный "
            f"Q_эфф ≥ {stats['q_foam_required_ls']:.1f} л/с и успешную ликвидацию."
        )
    disc.append(
        f"Средний индекс риска составил R̄ = {stats['risk_mean']:.3f}, "
        f"пиковое значение R_max = {stats['risk_max']:.3f}. "
        "Индекс является аддитивной комбинацией фазовой составляющей, "
        "наличия розлива, вторичных пожаров и относительной площади горения."
    )
    for txt in disc:
        _para_doc(txt)
        doc.add_paragraph()

    # Подписи к рисункам
    _heading("5. Подписи к рисункам и таблицам (шаблон)")
    _para_doc(
        "Рис. 1. Динамика ключевых показателей тушения пожара РВС: "
        "а — площадь зеркала горения S(t), м²; б — суммарный расход ОВ Q(t), л/с; "
        "в — число стволов охлаждения N(t); г — интегральный индекс риска R(t). "
        "Штриховые вертикальные линии — моменты пенных атак.",
        italic=True, size=9
    )
    doc.add_paragraph()
    _para_doc(
        "Рис. 2. Результаты работы Q-learning агента: "
        "а — Q-значения действий в финальном состоянии; "
        "б — распределение частоты выбора действий; "
        "в — накопленная функция награды R_sum(t).",
        italic=True, size=9
    )
    doc.add_paragraph()

    doc.add_paragraph()
    _para_doc(
        f"[Документ сформирован автоматически системой САУР ПСП v1.0 | "
        f"{stats['generated_at']}]",
        italic=True, size=8
    )

    doc.save(output_path)
    return output_path
